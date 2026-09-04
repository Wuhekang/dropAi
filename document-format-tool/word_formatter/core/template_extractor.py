from __future__ import annotations

from collections import Counter
from dataclasses import dataclass, field
from pathlib import Path
import re

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from word_formatter.core.word_converter import WordDocumentConverter
from word_formatter.core.rule_parser import NaturalLanguageRuleParser
from word_formatter.models.rules import DocumentRules, ParagraphRule, font_size_name_for_points


@dataclass(slots=True)
class TemplateExtractionResult:
    rules: DocumentRules
    notes: list[str] = field(default_factory=list)


class TemplateRuleExtractor:
    """从 Word 模板的真实示例内容中提取可确认的格式。

    样例段落优先于样式名称，因此学校模板即使使用“自定义样式 +
    直接格式”也可识别。当模板没有对应示例时，才回退到常见的 Word
    命名样式。
    """

    _FIGURE_CAPTION_RE = re.compile(
        r"^\s*(?:图|figure|fig\.?)\s*(?:[A-Za-z]?\d+|[一二三四五六七八九十百]+)"
        r"(?:[-—－–.．]\d+)*\s*[^\s]?",
        re.IGNORECASE,
    )
    _TABLE_CAPTION_RE = re.compile(
        r"^\s*(?:表|table)\s*(?:[A-Za-z]?\d+|[一二三四五六七八九十百]+)"
        r"(?:[-—－–.．]\d+)*\s*[^\s]?",
        re.IGNORECASE,
    )
    _REFERENCE_HEADING_RE = re.compile(
        r"^\s*(?:参考文献|主要参考文献|references?|bibliography)\s*[:：]?\s*$",
        re.IGNORECASE,
    )
    _REFERENCE_ITEM_RE = re.compile(
        r"^\s*(?:\[\s*\d+\s*\]|［\s*\d+\s*］|\d+\s*[.．](?!\d)\s*)\s*\S+"
    )
    _REFERENCE_END_RE = re.compile(r"^\s*(?:致谢|谢辞|附录|acknowledg(?:e)?ments?|appendix)\s*$", re.I)
    _COVER_METADATA_RE = re.compile(
        r"(?:学\s*院|院\s*[（(]?系|专\s*业|班\s*级|学\s*号|姓\s*名|学生姓名|"
        r"指导教师|学习中心|分\s*部|日\s*期|题\s*目)"
    )

    def __init__(self, converter: WordDocumentConverter | None = None) -> None:
        self.converter = converter or WordDocumentConverter()

    def extract(self, path: str | Path) -> TemplateExtractionResult:
        source = Path(path).expanduser()
        if source.suffix.lower() not in WordDocumentConverter.SUPPORTED_SUFFIXES or not source.is_file():
            raise ValueError("模板识别仅支持存在的 .doc、.docx 或 .dotx 文件")

        with self.converter.as_docx(source) as readable_path:
            document = Document(readable_path)
            rules = DocumentRules(name=f"从模板识别：{source.stem}")
            notes: list[str] = []
            self._extract_page(document, rules, notes)
            self._extract_page_number(document, rules, notes)
            samples = self._select_paragraph_samples(document)

            self._extract_sample_or_style(
                document,
                samples.get("normal_text"),
                rules.normal_text,
                (r"^Normal$", r"^(?:4)?正文$", r"^正文文本$", r"Body\s*Text"),
                "正文",
                notes,
            )
            for level in range(1, 5):
                heading_rule = getattr(rules, f"heading_{level}")
                self._extract_sample_or_style(
                    document,
                    samples.get(f"heading_{level}"),
                    heading_rule,
                    (
                        fr"Heading\s*{level}",
                        fr"标题\s*{level}",
                        fr"{level}\s*级标题",
                        fr"{'一二三四'[level - 1]}\s*级标题",
                        fr"(?:{level})?{'一二三四'[level - 1]}级目录",
                    ),
                    f"{level} 级标题",
                    notes,
                )
                if heading_rule.enabled:
                    # 语义层级比模板样式中可能错写的 outlineLvl 更可靠。
                    heading_rule.outline_level = level - 1
                    heading_rule.keep_with_next = True
            toc_title = next((p for p in document.paragraphs if re.fullmatch(r"(?:目\s*录|contents)", p.text.strip(), re.I)), None)
            self._extract_sample_or_style(document, toc_title, rules.toc_title, (r"TOC\s*Heading", r"目录标题"), "目录标题", notes)
            for level in range(1, 4):
                toc_sample = next((p for p in document.paragraphs if re.search(fr"(?:^|\s)(?:TOC|目录)\s*{level}", f"{p.style.style_id if p.style else ''} {p.style.name if p.style else ''}", re.I)), None)
                self._extract_sample_or_style(document, toc_sample, getattr(rules, f"toc_{level}"), (fr"TOC\s*{level}", fr"目录\s*{level}"), f"{level} 级目录", notes)
            self._extract_sample_or_style(
                document,
                samples.get("figure_caption"),
                rules.figure_caption,
                (r"Caption", r"题注", r"图题", r"图名"),
                "图名",
                notes,
            )
            self._extract_sample_or_style(
                document,
                samples.get("table_caption"),
                rules.table_caption,
                (r"表题", r"表名"),
                "表名",
                notes,
            )
            self._extract_sample_or_style(
                document,
                samples.get("reference"),
                rules.reference,
                (r"Bibliography", r"References?", r"参考文献"),
                "参考文献",
                notes,
            )

            written_specification = self._apply_written_specification(document, rules, notes)

            selected = None if written_specification else self._select_body_table(document)
            if selected is None:
                if not written_specification:
                    rules.table.enabled = False
                    notes.append("模板中没有可确认的正文表格，未覆盖当前表格规则。")
            else:
                table, table_index, selection_reason = selected
                self._extract_table(document, table, rules, notes)
                notes.append(f"表格样例选用模板中第 {table_index} 个表格：{selection_reason}。")
            return TemplateExtractionResult(rules, notes)

    @staticmethod
    def _apply_written_specification(
        document, rules: DocumentRules, notes: list[str]
    ) -> bool:
        """Read explicit format rules from a two-column specification table."""
        rows: dict[str, str] = {}
        for table in document.tables:
            for row in table.rows:
                if len(row.cells) < 2:
                    continue
                label = re.sub(r"\s+", "", row.cells[0].text)
                value = re.sub(r"\s+", " ", row.cells[1].text).strip()
                if label and value:
                    rows[label] = value
        required = {"页面设置", "目录", "正文", "图", "表"}
        if len(required & rows.keys()) < 3:
            return False

        parser = NaturalLanguageRuleParser()
        # A written specification is stronger evidence than incidental styles in
        # the explanatory document. Clear previously sampled paragraph values so
        # unrelated prose cannot leak into the returned rules.
        named_rules = [
            rules.normal_text,
            *(getattr(rules, f"heading_{level}") for level in range(1, 5)),
            rules.toc_title,
            *(getattr(rules, f"toc_{level}") for level in range(1, 4)),
            rules.figure_caption,
            rules.table_caption,
            rules.reference,
        ]
        for rule in named_rules:
            TemplateRuleExtractor._reset_rule_to_word_defaults(rule)
        known_fonts = [
            "宋体", "仿宋", "楷体", "黑体", "微软雅黑",
            "方正小标宋", "Times New Roman", "Arial",
        ]
        if "页面设置" in rows:
            normalized = re.sub(r"(?<=[页边距])为", "", rows["页面设置"])
            notes.extend(parser.apply(f"页面：{normalized}\n正文：{normalized}", rules))
        if "正文" in rows:
            notes.extend(parser.apply(rows["正文"], rules))
            for level in range(1, 5):
                heading = getattr(rules, f"heading_{level}")
                heading.outline_level = level - 1
                heading.keep_with_next = True
        if "图" in rows:
            parser._apply_figure_text(
                f"图名：{rows['图']}", rules.figure_caption, known_fonts, notes
            )
        if "表" in rows:
            caption_text = rows["表"].split("要求", 1)[0]
            if "居中" in rows["表"] and "居中" not in caption_text:
                caption_text += "，居中"
            parser._apply_figure_text(
                f"表名：{caption_text}", rules.table_caption, known_fonts, notes
            )
        if "目录" in rows:
            title_text, _, content_text = rows["目录"].partition("目录内容")
            parser._apply_heading_text(
                f"目录标题：{title_text}", rules.toc_title, 1, known_fonts, notes
            )
            rules.toc_title.outline_level = 9
            content_text = f"目录内容{content_text}" if content_text else rows["目录"]
            for level in range(1, 4):
                rule = getattr(rules, f"toc_{level}")
                parser._apply_heading_text(
                    content_text, rule, level, known_fonts, notes
                )
                rule.outline_level = 9
        if "参考文献" in rows:
            _, _, content_text = rows["参考文献"].partition("内容部分")
            parser._apply_heading_text(
                content_text or rows["参考文献"],
                rules.reference,
                1,
                known_fonts,
                notes,
            )
            rules.reference.outline_level = 9
        notes.append(
            "检测到撰写规范表：已直接提取页面、正文、标题、目录和图表题注规则，不复制规范说明正文。"
        )
        return True

    @staticmethod
    def _extract_page(document, rules: DocumentRules, notes: list[str]) -> None:
        # 只读取模板首节；提取器不展开或合并其他节，也不修改目标文档。
        section = document.sections[0]
        page = rules.page_setup
        page.enabled = True
        page.width_mm = section.page_width.mm
        page.height_mm = section.page_height.mm
        page.margin_top_mm = section.top_margin.mm
        page.margin_bottom_mm = section.bottom_margin.mm
        page.margin_left_mm = section.left_margin.mm
        page.margin_right_mm = section.right_margin.mm
        notes.append(
            "页面规则仅来自模板首节的纸张和页边距；模板其他节未被折叠到该规则中。"
        )

    @staticmethod
    def _extract_page_number(document, rules: DocumentRules, notes: list[str]) -> None:
        has_page_field = False
        for section in document.sections:
            for reference in section._sectPr.findall(qn("w:footerReference")):
                relationship_id = reference.get(qn("r:id"))
                part = document.part.related_parts.get(relationship_id)
                if part is None:
                    continue
                instructions = part.element.findall(".//" + qn("w:instrText"))
                if any(re.search(r"\bPAGE\b", item.text or "", re.I) for item in instructions):
                    has_page_field = True
                    break
            if has_page_field:
                break
        rules.page_number.enabled = has_page_field
        rules.page_number.settings = {
            "normalize_existing": has_page_field,
            "alignment": "center",
            "font_name": "Times New Roman",
            "font_size_pt": 10.5,
        }
        if has_page_field:
            notes.append("模板含页码域：处理时将保留前置页，并规范已有编号节为连续的居中页码。")

    def _extract_sample_or_style(
        self,
        document,
        sample: Paragraph | None,
        rule: ParagraphRule,
        fallback_patterns: tuple[str, ...],
        label: str,
        notes: list[str],
    ) -> None:
        if sample is not None:
            self._apply_paragraph_sample(document, sample, rule)
            preview = re.sub(r"\s+", " ", sample.text.strip())
            if len(preview) > 38:
                preview = f"{preview[:35]}…"
            notes.append(
                f"已从模板实际{label}段落“{preview}”合并样式继承和直接格式。"
            )
            return
        style = self._find_style(document, fallback_patterns)
        if style is not None:
            self._reset_rule_to_word_defaults(rule)
            self._apply_document_defaults(document, rule)
            self._apply_style(style, rule)
            rule.enabled = True
            notes.append(f"未找到实际{label}样例，已回退识别样式：{style.name}。")
        else:
            rule.enabled = False
            notes.append(f"模板中没有可确认的{label}样例，未覆盖该规则。")

    def _select_paragraph_samples(self, document) -> dict[str, Paragraph]:
        candidates: dict[str, list[tuple[int, int, Paragraph]]] = {
            "normal_text": [],
            "heading_1": [],
            "heading_2": [],
            "heading_3": [],
            "heading_4": [],
            "figure_caption": [],
            "table_caption": [],
            "reference": [],
        }
        in_references = False
        for index, paragraph in enumerate(document.paragraphs):
            text = re.sub(r"\s+", " ", paragraph.text).strip()
            if not text:
                continue
            if self._REFERENCE_HEADING_RE.match(text):
                in_references = True
                continue
            if in_references and self._REFERENCE_END_RE.match(text):
                in_references = False

            identities = self._style_identities(paragraph)
            if self._is_toc_style(identities) or text == "目录":
                continue
            if self._FIGURE_CAPTION_RE.match(text):
                candidates["figure_caption"].append((150, index, paragraph))
                continue
            if self._TABLE_CAPTION_RE.match(text):
                candidates["table_caption"].append((150, index, paragraph))
                continue

            # 参考文献常使用“1. 作者…”，必须在通用数字标题规则前判断。
            reference_score = 0
            if any(re.search(r"Bibliography|References?|参考文献", item, re.I) for item in identities):
                reference_score += 120
            if self._REFERENCE_ITEM_RE.match(text):
                reference_score += 90
            if in_references and len(text) >= 8:
                reference_score += 50
            if reference_score >= 90:
                candidates["reference"].append((reference_score, index, paragraph))
                continue

            style_level = self._style_heading_level(paragraph)
            outline_level = self._outline_heading_level(paragraph)
            numbered_level = self._numbered_heading_level(text)
            level = style_level or outline_level or numbered_level
            if level is not None and 1 <= level <= 4 and not self._REFERENCE_END_RE.match(text):
                score = 145 if style_level == level else 135 if outline_level == level else 105
                if len(text) <= 50:
                    score += 20
                elif len(text) > 100:
                    score -= 70
                if self._looks_bold(paragraph):
                    score += 10
                if text.endswith(("。", "；", ";")):
                    score -= 45
                if score >= 80:
                    candidates[f"heading_{level}"].append((score, index, paragraph))
                    continue

            if len(text) < 4 or len(text) > 2000:
                continue
            body_score = 10
            if any(re.search(r"^(?:Normal|4?正文|正文文本|Body\s*Text)$", item, re.I) for item in identities):
                body_score += 50
            if 20 <= len(text) <= 600:
                body_score += 30
            elif len(text) < 8:
                body_score -= 10
            if text.endswith(("。", ".", "；", ";")):
                body_score += 12
            if self._COVER_METADATA_RE.search(text) and len(text) < 45:
                body_score -= 70
            if paragraph.alignment is not None and int(paragraph.alignment) == 1:
                body_score -= 12
            if body_score > 0:
                candidates["normal_text"].append((body_score, index, paragraph))

        selected: dict[str, Paragraph] = {}
        for role, items in candidates.items():
            chosen = self._choose_representative(document, items)
            if chosen is not None:
                selected[role] = chosen
        return selected

    def _choose_representative(
        self, document, candidates: list[tuple[int, int, Paragraph]]
    ) -> Paragraph | None:
        if not candidates:
            return None
        groups: dict[tuple, list[tuple[int, int, Paragraph]]] = {}
        for item in candidates:
            signature = self._paragraph_signature(document, item[2])
            groups.setdefault(signature, []).append(item)
        best_group = max(
            groups.values(),
            key=lambda group: (
                sum(item[0] for item in group) + min(len(group), 12) * 24,
                max(item[0] for item in group),
                len(group),
            ),
        )
        return max(best_group, key=lambda item: (item[0], len(item[2].text), -item[1]))[2]

    def _paragraph_signature(self, document, paragraph: Paragraph) -> tuple:
        rule = ParagraphRule()
        self._apply_paragraph_sample(document, paragraph, rule)
        return (
            rule.chinese_font,
            rule.latin_font,
            round(rule.font_size_pt, 2),
            rule.bold,
            rule.italic,
            rule.underline,
            rule.alignment,
            rule.line_spacing_mode,
            round(rule.fixed_line_spacing_pt, 2),
            round(rule.multiple_line_spacing, 2),
            rule.special_indent_mode,
            round(rule.special_indent_chars, 2),
        )

    @classmethod
    def _style_heading_level(cls, paragraph: Paragraph) -> int | None:
        for identity in cls._style_identities(paragraph):
            compact = re.sub(r"[\s_\-]", "", identity)
            match = re.search(r"(?:Heading|标题)([1-4])", compact, re.I)
            if match:
                return int(match.group(1))
            match = re.search(r"([1-4])级标题", compact)
            if match:
                return int(match.group(1))
            for level, chinese in enumerate("一二三四", start=1):
                if f"{chinese}级标题" in compact or re.fullmatch(
                    fr"(?:{level})?{chinese}级目录", compact
                ):
                    return level
        return None

    @classmethod
    def _outline_heading_level(cls, paragraph: Paragraph) -> int | None:
        p_pr = paragraph._p.pPr
        if p_pr is not None:
            outline = p_pr.find(qn("w:outlineLvl"))
            if outline is not None:
                value = cls._safe_int(outline.get(qn("w:val")))
                if value is not None and 0 <= value <= 3:
                    return value + 1
        style = paragraph.style
        seen: set[int] = set()
        while style is not None and id(style) not in seen:
            seen.add(id(style))
            style_p_pr = style._element.pPr
            if style_p_pr is not None:
                outline = style_p_pr.find(qn("w:outlineLvl"))
                if outline is not None:
                    value = cls._safe_int(outline.get(qn("w:val")))
                    if value is not None and 0 <= value <= 3:
                        return value + 1
            style = style.base_style
        return None

    @staticmethod
    def _numbered_heading_level(text: str) -> int | None:
        if re.match(r"^\s*第\s*[一二三四五六七八九十百\d]+\s*章(?:\s|[:：]|$)", text):
            return 1
        if re.match(r"^\s*[一二三四五六七八九十百]+\s*[、.．]\s*\S+", text):
            return 1
        if re.match(r"^\s*[（(][一二三四五六七八九十百]+[）)]\s*\S+", text):
            return 2
        if re.match(r"^\s*[（(]\d+[）)]\s*\S+", text):
            return 4
        match = re.match(r"^\s*\d+(?:[.．]\d+){0,3}", text)
        if match:
            prefix = match.group(0).strip()
            end = match.end()
            if end < len(text) and text[end].isdigit():
                return None
            level = len(re.findall(r"[.．]", prefix)) + 1
            if 1 <= level <= 4:
                return level
        if re.match(r"^\s*\d+\s*、\s*\S+", text):
            return 1
        return None

    @staticmethod
    def _style_identities(paragraph: Paragraph) -> tuple[str, ...]:
        style = paragraph.style
        if style is None:
            return ()
        return tuple(item for item in (style.style_id or "", style.name or "") if item)

    @staticmethod
    def _is_toc_style(identities: tuple[str, ...]) -> bool:
        # 仅跳过 Word 真正的 TOC/目录条目样式；学校模板常把
        # 正文标题样式命名为“1一级目录/2二级目录”，不能误过滤。
        return any(
            re.fullmatch(r"\s*(?:TOC\s*\d+|目录\s*\d*|目录文本)\s*", item, re.I)
            is not None
            for item in identities
        )

    @staticmethod
    def _looks_bold(paragraph: Paragraph) -> bool:
        visible = [run for run in paragraph.runs if run.text.strip()]
        if visible and any(run.bold is True for run in visible):
            return True
        style = paragraph.style
        return bool(style is not None and style.font.bold is True)

    @staticmethod
    def _find_style(document, patterns: tuple[str, ...]):
        for style in document.styles:
            identities = (style.style_id or "", style.name or "")
            if any(
                re.search(pattern, identity, re.IGNORECASE)
                for pattern in patterns
                for identity in identities
            ):
                return style
        return None

    @classmethod
    def _apply_paragraph_sample(
        cls,
        document,
        paragraph: Paragraph,
        rule: ParagraphRule,
        table: Table | None = None,
    ) -> None:
        rule.enabled = True
        cls._reset_rule_to_word_defaults(rule)
        cls._apply_document_defaults(document, rule)
        if table is not None:
            cls._apply_table_style(table, rule)
        if paragraph.style is not None:
            cls._apply_style(paragraph.style, rule)
        p_pr = paragraph._p.pPr
        if p_pr is not None:
            cls._apply_ppr(p_pr, rule)
            paragraph_mark_r_pr = p_pr.find(qn("w:rPr"))
            if paragraph_mark_r_pr is not None:
                cls._apply_rpr(paragraph_mark_r_pr, rule)
        visible_runs = [run for run in paragraph.runs if run.text.strip()]
        if visible_runs:
            dominant_run = max(visible_runs, key=lambda run: len(run.text.strip()))
            if dominant_run._element.rPr is not None:
                cls._apply_rpr(dominant_run._element.rPr, rule)

    @staticmethod
    def _reset_rule_to_word_defaults(rule: ParagraphRule) -> None:
        """清除 DocumentRules 的预设值，避免把“未声明”误当成模板证据。"""

        rule.bold = False
        rule.italic = False
        rule.underline = False
        rule.character_spacing_mode = "standard"
        rule.character_spacing_pt = 0.0
        rule.direction = "ltr"
        rule.outline_level = 9
        rule.left_indent_cm = 0.0
        rule.right_indent_cm = 0.0
        rule.left_indent_chars = 0.0
        rule.right_indent_chars = 0.0
        rule.special_indent_mode = "none"
        rule.special_indent_chars = 0.0
        rule.first_line_indent_chars = 0.0
        rule.line_spacing_mode = "single"
        rule.space_before_unit = "pt"
        rule.space_after_unit = "pt"
        rule.space_before_pt = 0.0
        rule.space_after_pt = 0.0
        rule.space_before_lines = 0.0
        rule.space_after_lines = 0.0
        rule.alignment = "left"
        rule.snap_to_grid = False
        rule.auto_adjust_right_indent = True
        rule.widow_control = True
        rule.keep_with_next = False
        rule.keep_lines_together = False
        rule.page_break_before = False

    @classmethod
    def _apply_document_defaults(cls, document, rule: ParagraphRule) -> None:
        defaults = document.styles.element.find(qn("w:docDefaults"))
        if defaults is None:
            return
        r_pr = defaults.find(f"{qn('w:rPrDefault')}/{qn('w:rPr')}")
        if r_pr is not None:
            cls._apply_rpr(r_pr, rule)
        p_pr = defaults.find(f"{qn('w:pPrDefault')}/{qn('w:pPr')}")
        if p_pr is not None:
            cls._apply_ppr(p_pr, rule)

    @classmethod
    def _apply_style(cls, style, rule: ParagraphRule) -> None:
        rule.enabled = True
        chain = []
        current = style
        seen: set[int] = set()
        while current is not None and id(current) not in seen:
            seen.add(id(current))
            chain.append(current)
            current = current.base_style
        for candidate in reversed(chain):
            r_pr = candidate._element.rPr
            if r_pr is not None:
                cls._apply_rpr(r_pr, rule)
            p_pr = candidate._element.pPr
            if p_pr is not None:
                cls._apply_ppr(p_pr, rule)

    @classmethod
    def _apply_table_style(cls, table: Table, rule: ParagraphRule) -> None:
        style = table.style
        chain = []
        seen: set[int] = set()
        while style is not None and id(style) not in seen:
            seen.add(id(style))
            chain.append(style)
            style = style.base_style
        for candidate in reversed(chain):
            r_pr = candidate._element.find(qn("w:rPr"))
            if r_pr is not None:
                cls._apply_rpr(r_pr, rule)
            p_pr = candidate._element.find(qn("w:pPr"))
            if p_pr is not None:
                cls._apply_ppr(p_pr, rule)
            for conditional in candidate._element.findall(qn("w:tblStylePr")):
                if conditional.get(qn("w:type")) != "wholeTable":
                    continue
                r_pr = conditional.find(qn("w:rPr"))
                if r_pr is not None:
                    cls._apply_rpr(r_pr, rule)
                p_pr = conditional.find(qn("w:pPr"))
                if p_pr is not None:
                    cls._apply_ppr(p_pr, rule)

    @classmethod
    def _apply_rpr(cls, r_pr, rule: ParagraphRule) -> None:
        fonts = r_pr.find(qn("w:rFonts"))
        if fonts is not None:
            east_asia = fonts.get(qn("w:eastAsia"))
            latin = fonts.get(qn("w:ascii")) or fonts.get(qn("w:hAnsi"))
            if east_asia:
                rule.chinese_font = east_asia
            if latin:
                rule.latin_font = latin
                rule.number_font = latin
        size = r_pr.find(qn("w:sz"))
        if size is None:
            size = r_pr.find(qn("w:szCs"))
        if size is not None:
            value = cls._safe_float(size.get(qn("w:val")))
            if value is not None and value > 0:
                rule.font_size_pt = value / 2
                rule.font_size_name = font_size_name_for_points(rule.font_size_pt)
        for tag, attribute in (("b", "bold"), ("i", "italic")):
            element = r_pr.find(qn(f"w:{tag}"))
            if element is not None:
                setattr(rule, attribute, cls._on_off(element))
        underline = r_pr.find(qn("w:u"))
        if underline is not None:
            rule.underline = underline.get(qn("w:val"), "single") not in {"none", "0", "false", "off"}
        spacing = r_pr.find(qn("w:spacing"))
        if spacing is not None:
            value = cls._safe_float(spacing.get(qn("w:val")))
            if value is not None:
                points = value / 20
                rule.character_spacing_mode = "expanded" if points > 0 else "condensed" if points < 0 else "standard"
                rule.character_spacing_pt = abs(points)

    @classmethod
    def _apply_ppr(cls, p_pr, rule: ParagraphRule) -> None:
        jc = p_pr.find(qn("w:jc"))
        if jc is not None:
            rule.alignment = {
                "left": "left",
                "center": "center",
                "right": "right",
                "both": "justify",
                "distribute": "justify",
            }.get(jc.get(qn("w:val")), rule.alignment)
        spacing = p_pr.find(qn("w:spacing"))
        if spacing is not None:
            before_lines = cls._safe_float(spacing.get(qn("w:beforeLines")))
            before = cls._safe_float(spacing.get(qn("w:before")))
            after_lines = cls._safe_float(spacing.get(qn("w:afterLines")))
            after = cls._safe_float(spacing.get(qn("w:after")))
            if before_lines is not None:
                rule.space_before_unit = "line"
                rule.space_before_lines = before_lines / 100
            elif before is not None:
                rule.space_before_unit = "pt"
                rule.space_before_pt = before / 20
            if after_lines is not None:
                rule.space_after_unit = "line"
                rule.space_after_lines = after_lines / 100
            elif after is not None:
                rule.space_after_unit = "pt"
                rule.space_after_pt = after / 20
            line = cls._safe_float(spacing.get(qn("w:line")))
            line_rule = spacing.get(qn("w:lineRule"), "auto")
            if line is not None:
                if line_rule == "exact":
                    rule.line_spacing_mode = "fixed"
                    rule.fixed_line_spacing_pt = line / 20
                elif line_rule == "atLeast":
                    rule.line_spacing_mode = "at_least"
                    rule.minimum_line_spacing_pt = line / 20
                else:
                    multiple = line / 240
                    if abs(multiple - 1) < 0.01:
                        rule.line_spacing_mode = "single"
                    elif abs(multiple - 1.5) < 0.01:
                        rule.line_spacing_mode = "1.5"
                    elif abs(multiple - 2) < 0.01:
                        rule.line_spacing_mode = "double"
                    else:
                        rule.line_spacing_mode = "multiple"
                        rule.multiple_line_spacing = multiple
        ind = p_pr.find(qn("w:ind"))
        if ind is not None:
            left = cls._safe_float(ind.get(qn("w:left")))
            right = cls._safe_float(ind.get(qn("w:right")))
            if left is not None:
                rule.left_indent_cm = left / 567
            if right is not None:
                rule.right_indent_cm = right / 567
            left_chars = cls._safe_float(ind.get(qn("w:leftChars")))
            right_chars = cls._safe_float(ind.get(qn("w:rightChars")))
            if left_chars is not None:
                rule.left_indent_chars = left_chars / 100
            if right_chars is not None:
                rule.right_indent_chars = right_chars / 100
            first_line_chars = cls._safe_float(ind.get(qn("w:firstLineChars")))
            hanging_chars = cls._safe_float(ind.get(qn("w:hangingChars")))
            first_line = cls._safe_float(ind.get(qn("w:firstLine")))
            hanging = cls._safe_float(ind.get(qn("w:hanging")))
            if first_line_chars is not None:
                chars = first_line_chars / 100
                rule.special_indent_mode = "first_line" if chars else "none"
                rule.special_indent_chars = chars
                rule.first_line_indent_chars = chars
            elif first_line is not None:
                chars = first_line / max(rule.font_size_pt * 20, 1)
                rule.special_indent_mode = "first_line" if chars else "none"
                rule.special_indent_chars = chars
                rule.first_line_indent_chars = chars
            elif hanging_chars is not None:
                chars = hanging_chars / 100
                rule.special_indent_mode = "hanging" if chars else "none"
                rule.special_indent_chars = chars
                rule.first_line_indent_chars = 0
            elif hanging is not None:
                chars = hanging / max(rule.font_size_pt * 20, 1)
                rule.special_indent_mode = "hanging" if chars else "none"
                rule.special_indent_chars = chars
                rule.first_line_indent_chars = 0

        outline = p_pr.find(qn("w:outlineLvl"))
        if outline is not None:
            value = cls._safe_int(outline.get(qn("w:val")))
            if value is not None:
                rule.outline_level = max(0, min(value, 9))
        boolean_properties = {
            "snapToGrid": "snap_to_grid",
            "adjustRightInd": "auto_adjust_right_indent",
            "widowControl": "widow_control",
            "keepNext": "keep_with_next",
            "keepLines": "keep_lines_together",
            "pageBreakBefore": "page_break_before",
        }
        for tag, attribute in boolean_properties.items():
            element = p_pr.find(qn(f"w:{tag}"))
            if element is not None:
                setattr(rule, attribute, cls._on_off(element))
        bidi = p_pr.find(qn("w:bidi"))
        if bidi is not None:
            rule.direction = "rtl" if cls._on_off(bidi) else "ltr"

    @classmethod
    def _select_body_table(cls, document) -> tuple[Table, int, str] | None:
        blocks = list(cls._iter_body_blocks(document))
        table_indexes = {id(table._tbl): index for index, table in enumerate(document.tables, start=1)}
        scored: list[tuple[int, int, Table, str]] = []
        for position, block in enumerate(blocks):
            if not isinstance(block, Table):
                continue
            table_index = table_indexes.get(id(block._tbl), len(scored) + 1)
            previous = cls._nearby_paragraph_text(blocks, position, -1)
            following = cls._nearby_paragraph_text(blocks, position, 1)
            previous_is_caption = bool(previous and cls._TABLE_CAPTION_RE.match(previous))
            following_is_caption = bool(following and cls._TABLE_CAPTION_RE.match(following))
            rows = len(block.rows)
            columns = len(block.columns) if rows else 0
            texts = [
                re.sub(r"\s+", " ", cell.text).strip()
                for row in block.rows
                for cell in row.cells
                if cell.text.strip()
            ]
            combined = " ".join(texts)
            metadata_hits = len(cls._COVER_METADATA_RE.findall(combined))
            cover_metadata = metadata_hits >= 2 and rows <= 12 and len(combined) < 500
            score = 0
            reasons = []
            if previous_is_caption:
                score += 260
                reasons.append("紧跟在表名之后")
            elif following_is_caption:
                score += 100
                reasons.append("邻近表名")
            if rows >= 2 and columns >= 2:
                score += 55
                reasons.append("具有正文数据网格")
            if len(texts) >= 4:
                score += min(len(texts), 20)
            if rows <= 1 or columns <= 1:
                score -= 45
            if position < 12 and not previous_is_caption:
                score -= 35
            if cover_metadata:
                score -= 220
                reasons.append("封面元数据特征已降权")
            scored.append((score, -table_index, block, "、".join(reasons) or "正文区域的最佳表格候选"))
        if not scored:
            return None
        score, negative_index, table, reason = max(scored, key=lambda item: (item[0], item[1]))
        if score < -100:
            return None
        return table, -negative_index, reason

    @staticmethod
    def _iter_body_blocks(document):
        body = document.element.body
        for child in body.iterchildren():
            if child.tag == qn("w:p"):
                yield Paragraph(child, document)
            elif child.tag == qn("w:tbl"):
                yield Table(child, document)

    @staticmethod
    def _nearby_paragraph_text(blocks: list, position: int, direction: int) -> str:
        cursor = position + direction
        inspected = 0
        while 0 <= cursor < len(blocks) and inspected < 3:
            block = blocks[cursor]
            if isinstance(block, Table):
                break
            text = re.sub(r"\s+", " ", block.text).strip()
            if text:
                return text
            cursor += direction
            inspected += 1
        return ""

    @classmethod
    def _extract_table(cls, document, table: Table, rules: DocumentRules, notes: list[str]) -> None:
        rule = rules.table
        rule.enabled = True
        # 0 表示保留目标表格自身几何，不把单个样例的行高/列宽强加给所有表格。
        rule.row_height_mm = 0.0
        rule.column_width_mm = 0.0

        borders = cls._effective_table_borders(table)
        cls._apply_border_evidence(table, borders, rule, notes)
        paragraph = cls._representative_table_paragraph(table)
        if paragraph is not None:
            cls._apply_paragraph_sample(document, paragraph, rule, table=table)
            notes.append("已从所选正文表格的非空单元格合并表格样式、段落样式和直接格式。")
        else:
            notes.append("所选正文表格没有可用的非空单元格，未提取字体和段落证据。")
        header_paragraphs = [
            paragraph
            for cell in (table.rows[0].cells if table.rows else [])
            for paragraph in cell.paragraphs
            if paragraph.text.strip()
        ]
        if header_paragraphs:
            bold_count = sum(cls._looks_bold(paragraph) for paragraph in header_paragraphs)
            rule.header_row_bold = bold_count * 2 >= len(header_paragraphs)
            if rule.header_row_bold:
                notes.append("所选正文表格首行以粗体为主，目标表格首行将统一加粗。")
        # 长表跨页时重复表头是稳定的论文排版要求；即使单页模板样例没有
        # 显式 tblHeader，也保留默认开启，避免续页失去列名。
        rule.repeat_header_row = True

    @classmethod
    def _effective_table_borders(cls, table: Table) -> dict[str, dict[str, str | None]]:
        edges: dict[str, dict[str, str | None]] = {}
        style = table.style
        chain = []
        seen: set[int] = set()
        while style is not None and id(style) not in seen:
            seen.add(id(style))
            chain.append(style)
            style = style.base_style
        for candidate in reversed(chain):
            tbl_pr = candidate._element.find(qn("w:tblPr"))
            if tbl_pr is not None:
                cls._merge_border_container(tbl_pr.find(qn("w:tblBorders")), edges)
            for conditional in candidate._element.findall(qn("w:tblStylePr")):
                if conditional.get(qn("w:type")) == "wholeTable":
                    tbl_pr = conditional.find(qn("w:tblPr"))
                    if tbl_pr is not None:
                        cls._merge_border_container(tbl_pr.find(qn("w:tblBorders")), edges)
        cls._merge_border_container(table._tbl.tblPr.find(qn("w:tblBorders")), edges)
        if not edges:
            observations: dict[str, list[dict[str, str | None]]] = {}
            for row in table.rows:
                for cell in row.cells:
                    tc_pr = cell._tc.tcPr
                    if tc_pr is None:
                        continue
                    container = tc_pr.find(qn("w:tcBorders"))
                    if container is None:
                        continue
                    for name in ("top", "bottom", "left", "right", "insideH", "insideV"):
                        element = container.find(qn(f"w:{name}"))
                        if element is not None:
                            observations.setdefault(name, []).append(cls._border_data(element))
            for name, values in observations.items():
                val, size, color = Counter(
                    (item.get("val"), item.get("sz"), item.get("color")) for item in values
                ).most_common(1)[0][0]
                edges[name] = {"val": val, "sz": size, "color": color}
        return edges

    @classmethod
    def _merge_border_container(cls, container, target: dict[str, dict[str, str | None]]) -> None:
        if container is None:
            return
        for name in ("top", "bottom", "left", "right", "insideH", "insideV"):
            element = container.find(qn(f"w:{name}"))
            if element is not None:
                target[name] = cls._border_data(element)

    @staticmethod
    def _border_data(element) -> dict[str, str | None]:
        return {
            "val": element.get(qn("w:val")),
            "sz": element.get(qn("w:sz")),
            "color": element.get(qn("w:color")),
        }

    @classmethod
    def _apply_border_evidence(cls, table, borders, rule, notes: list[str]) -> None:
        def active(name: str) -> bool:
            return borders.get(name, {}).get("val") not in {None, "nil", "none", "0"}

        if not borders or not any(active(name) for name in borders):
            rule.border_style = "none"
            notes.append("所选表格及其表格样式没有有效框线定义，识别为无框线。")
            return
        if (
            active("top")
            and active("bottom")
            and not active("left")
            and not active("right")
            and not active("insideV")
        ):
            rule.border_style = "three_line"
        else:
            rule.border_style = "grid"

        outer_widths = [
            cls._border_width(borders[name])
            for name in ("top", "bottom", "left", "right")
            if active(name) and cls._border_width(borders[name]) is not None
        ]
        inner_widths = [
            cls._border_width(borders[name])
            for name in ("insideH", "insideV")
            if active(name) and cls._border_width(borders[name]) is not None
        ]
        if outer_widths:
            rule.outer_border_width_pt = sum(outer_widths) / len(outer_widths)
        if inner_widths:
            rule.inner_border_width_pt = sum(inner_widths) / len(inner_widths)
        elif outer_widths:
            rule.inner_border_width_pt = min(outer_widths)
        colors = [
            data.get("color")
            for data in borders.values()
            if data.get("color") not in {None, "auto"}
        ]
        if colors:
            rule.border_color = str(Counter(colors).most_common(1)[0][0]).lstrip("#").upper()
        label = "三线表" if rule.border_style == "three_line" else "全框线"
        notes.append(f"已从所选表格的有效框线证据识别为{label}。")

    @staticmethod
    def _border_width(data: dict[str, str | None]) -> float | None:
        try:
            return float(data.get("sz")) / 8 if data.get("sz") is not None else None
        except (TypeError, ValueError):
            return None

    @staticmethod
    def _representative_table_paragraph(table: Table) -> Paragraph | None:
        candidates: list[tuple[tuple[str, int | None], int, int, Paragraph]] = []
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = re.sub(r"\s+", " ", paragraph.text).strip()
                    if not text:
                        continue
                    style = paragraph.style
                    signature = (
                        (style.style_id if style else "") or (style.name if style else ""),
                        int(paragraph.alignment) if paragraph.alignment is not None else None,
                    )
                    score = (45 if row_index > 0 else 15) + min(len(text), 80)
                    candidates.append((signature, score, len(text), paragraph))
        if not candidates:
            return None
        # 表格中的少数长说明单元格常有直接“两端对齐”，不能让它覆盖
        # 大多数普通单元格的居中样式；先取最常见格式签名，再选其代表段落。
        counts = Counter(item[0] for item in candidates)
        signature = max(counts, key=lambda item: counts[item])
        matching = [item for item in candidates if item[0] == signature]
        return max(matching, key=lambda item: (item[1], item[2]))[3]

    @staticmethod
    def _safe_int(value: str | None) -> int | None:
        try:
            return int(value) if value is not None else None
        except (TypeError, ValueError):
            return None

    @staticmethod
    def _safe_float(value: str | None) -> float | None:
        try:
            return float(value) if value is not None else None
        except (TypeError, ValueError):
            return None

    @staticmethod
    def _on_off(element) -> bool:
        return element.get(qn("w:val"), "1") not in {"0", "false", "off", "none"}
