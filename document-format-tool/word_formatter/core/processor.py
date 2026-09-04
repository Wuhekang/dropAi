from __future__ import annotations

from pathlib import Path
from copy import deepcopy
import os
import re
import tempfile
import time

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import (
    WD_CELL_VERTICAL_ALIGNMENT,
    WD_ROW_HEIGHT_RULE,
    WD_TABLE_ALIGNMENT,
)
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Cm, Mm, Pt
from docx.table import Table
from docxcompose.composer import Composer

from word_formatter.core.analyzer import DocumentAnalyzer
from word_formatter.core.finalizer import finalize_docx, is_red_font_value
from word_formatter.core.word_converter import WordDocumentConverter
from word_formatter.models.results import ChangeRecord, ProcessResult
from word_formatter.models.rules import (
    DocumentRules,
    ParagraphRule,
    TableRule,
    enforce_locked_table_policy,
)


ALIGNMENTS = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}
VERTICAL_ALIGNMENTS = {
    "top": WD_CELL_VERTICAL_ALIGNMENT.TOP,
    "center": WD_CELL_VERTICAL_ALIGNMENT.CENTER,
    "bottom": WD_CELL_VERTICAL_ALIGNMENT.BOTTOM,
}


class DocumentProcessor:
    """安全处理 DOCX：只保存到新文件，绝不覆盖原文件。"""

    def process(
        self,
        source_path: str | Path,
        rules: DocumentRules,
        output_path: str | Path | None = None,
        template_path: str | Path | None = None,
    ) -> ProcessResult:
        source = Path(source_path).resolve()
        self._validate_source(source)
        output = Path(output_path).resolve() if output_path else self.default_output_path(source)
        if output == source:
            raise ValueError("输出文件不能与原文件相同")
        output.parent.mkdir(parents=True, exist_ok=True)
        result = ProcessResult(source_path=source, output_path=output)
        try:
            # CLI callers already apply this policy before reporting their rule
            # summary. Enforce it again here so direct/library callers cannot
            # bypass the fixed table contract.
            enforce_locked_table_policy(rules)
            document = self._compose_with_template_front(source, Path(template_path).resolve(), result) if template_path else Document(source)
            content_start = self._main_content_start(document)
            content_start += self._ensure_toc(document, rules, result, content_start)
            self._enforce_global_paragraph_policy(document, result)
            if content_start > 1:
                result.warnings.append(
                    f"已保留正文起点之前的 {content_start - 1} 个段落（封面、声明或目录），不套用正文格式。"
                )
            if rules.page_setup.enabled:
                self._apply_page_setup(document, rules, result)
            reference_paragraphs = self._reference_paragraphs(document, content_start)
            self._apply_toc(document, rules, result)
            if rules.figure_caption.enabled:
                self._apply_figure_captions(
                    document, rules.figure_caption, result, content_start
                )
            if rules.table_caption.enabled:
                self._apply_table_captions(
                    document, rules.table_caption, result, content_start
                )
            if rules.reference.enabled:
                self._apply_references(reference_paragraphs, rules.reference, result)
            if rules.normal_text.enabled:
                self._apply_normal_text(
                    document,
                    rules.normal_text,
                    result,
                    excluded_elements={id(paragraph._p) for _, paragraph in reference_paragraphs},
                    start_index=content_start,
                )
            self._apply_headings(document, rules, result, content_start)
            self._start_chapters_on_new_pages(document, content_start, result)
            self._exclude_non_content_toc_entries(document, content_start)
            if rules.table.enabled:
                self._apply_tables(document, rules.table, result, content_start)
            if rules.page_number.enabled:
                self._apply_page_numbers(document, rules.page_number.settings, result)
            if rules.normal_text.number_font != rules.normal_text.latin_font:
                result.warnings.append(
                    "python-docx 无法在不拆分文本运行块的情况下区分英文与数字字体；第一版数字字体暂按英文字体处理。"
                )
            if rules.table.number_font != rules.table.latin_font:
                result.warnings.append(
                    "表格中的数字字体暂按表格英文字体处理，以避免拆分文字运行块。"
                )
            self._request_field_update(document, result)
            document.save(output)
            # Remove review-only package parts before Word opens the file.
            # Some school templates contain stale comment extensions that make
            # Word reject field updates until those parts are stripped.
            cleanup = finalize_docx(output)
            if os.name == "nt":
                refresh_error = None
                for attempt in range(2):
                    try:
                        WordDocumentConverter().update_fields_in_place(output)
                        refreshed = Document(output)
                        self._apply_toc(refreshed, rules, result)
                        self._request_field_update(refreshed, result)
                        refreshed.save(output)
                        WordDocumentConverter().update_fields_in_place(output)
                        refresh_error = None
                        break
                    except Exception as exc:
                        refresh_error = exc
                        if attempt == 0:
                            time.sleep(0.75)
                if refresh_error is not None:
                    result.warnings.append(f"目录域已插入，但自动刷新失败：{refresh_error}")
            final_cleanup = finalize_docx(output)
            for key, value in final_cleanup.items():
                cleanup[key] += value
            if cleanup["comment_markup_removed"] or cleanup["comment_parts_removed"]:
                result.records.append(ChangeRecord(None, "审阅批注", "模板或原稿含批注", "全部移除", "最终稿固定规则"))
            if cleanup["red_fonts_blackened"]:
                result.records.append(ChangeRecord(None, "红色字体", "原稿残留红色直接格式", "统一改为黑色", "最终稿固定规则"))
            result.save_log(output.with_suffix(".log.json"))
            return result
        except Exception as exc:
            detail = str(exc).strip() or repr(exc)
            result.warnings.append(f"处理失败（{exc.__class__.__name__}）：{detail}")
            result.save_log(output.with_suffix(".failed.log.json"))
            raise RuntimeError(
                f"文档处理失败，原文件未改动。根因：{exc.__class__.__name__}: {detail}"
            ) from exc

    @classmethod
    def _compose_with_template_front(cls, source: Path, template: Path, result: ProcessResult):
        # Template extraction already supports legacy .doc/.dotx, but the
        # formatting phase must convert it as well. Passing an OLE .doc path
        # directly to python-docx produces a misleading missing-officeDocument
        # relationship error after the user confirms the rules.
        with WordDocumentConverter().as_docx(template) as readable_template:
            template_doc = Document(readable_template)
            source_doc = Document(source)
            if cls._looks_like_format_specification(template_doc):
                result.warnings.append(
                    "上传文件属于撰写/排版规范说明，仅提取格式规则，不复制其说明文字作为论文封面。"
                )
                return source_doc
            template_start = cls._main_content_start(template_doc)
            source_start = cls._main_content_start(source_doc)
            if template_start <= 1:
                result.warnings.append("模板未识别到独立前置页，保留论文原有前置内容。")
                return source_doc
            removed_review_text = cls._remove_template_review_artifacts(template_doc, template_start)
            if removed_review_text:
                result.records.append(ChangeRecord(None, "模板红字说明", f"{removed_review_text} 个红色文字节点", "未复制到最终稿", "最终稿固定规则"))
            added_sections = cls._isolate_front_matter_pages(template_doc, template_start)
            if added_sections:
                result.records.append(ChangeRecord(None, "前置页分节", "分页符或段前分页", f"新增 {added_sections} 个下一页分节符", "每个前置页独立成节"))
            cls._repair_missing_numbering_part(source_doc, result)
            cls._trim_body(template_doc, keep_before=template_start)
            cls._trim_body(source_doc, remove_before=source_start)
            Composer(template_doc).append(source_doc)
            result.records.append(ChangeRecord(None, "模板前置内容", f"论文原前置段落 {max(0, source_start - 1)} 个", f"复制模板前置段落 {template_start - 1} 个", "固定系统规则"))
            return template_doc

    @staticmethod
    def _remove_template_review_artifacts(document, content_start: int) -> int:
        removed = 0
        body = document.element.body
        paragraph_number = 0
        in_template_toc = False
        instruction_pattern = re.compile(
            r"(?:编写说明|提交时.{0,8}删除|请删除.{0,8}(?:提示|表格)|此页页码|白页.*偶数)",
            re.I,
        )
        for child in list(body.iterchildren()):
            if child.tag == qn("w:p"):
                paragraph_number += 1
            if paragraph_number >= content_start:
                break
            text = "".join(node.text or "" for node in child.xpath(".//w:t"))
            if child.tag == qn("w:p") and re.fullmatch(r"\s*目\s*录\s*", text):
                in_template_toc = True
            if in_template_toc:
                if child.tag == qn("w:p"):
                    # The template's sample TOC often ends in its own section.
                    # Once its visible contents are discarded, retaining that
                    # section creates a completely blank page before the real
                    # generated TOC.  The preceding abstract/front section is
                    # already the required next-page boundary.
                    p_pr = child.find(qn("w:pPr"))
                    if p_pr is not None:
                        section = p_pr.find(qn("w:sectPr"))
                        if section is not None:
                            p_pr.remove(section)
                        paragraph_mark = p_pr.find(qn("w:rPr"))
                        if paragraph_mark is None:
                            paragraph_mark = OxmlElement("w:rPr")
                            p_pr.append(paragraph_mark)
                        if paragraph_mark.find(qn("w:vanish")) is None:
                            paragraph_mark.append(OxmlElement("w:vanish"))
                    for node in list(child):
                        if node.tag != qn("w:pPr"):
                            child.remove(node)
                            removed += 1
                else:
                    body.remove(child)
                    removed += max(1, len(child.xpath(".//w:t")))
                continue
            border_colors = [
                element.get(qn("w:color"))
                for element in child.xpath(".//w:tblBorders/* | .//w:tcBorders/* | .//w:pBdr/*")
            ]
            review_container = (
                child.tag == qn("w:tbl")
                and (
                    any(is_red_font_value(value) for value in border_colors)
                    or bool(instruction_pattern.search(text))
                )
            )
            if review_container:
                body.remove(child)
                removed += max(1, len(child.xpath(".//w:t")))
                continue
            for run in list(child.xpath(".//w:r")):
                colors = run.xpath("./w:rPr/w:color")
                if not any(is_red_font_value(color.get(qn("w:val"))) for color in colors):
                    continue
                for text_node in list(run.xpath(".//w:t | .//w:delText")):
                    parent = text_node.getparent()
                    if parent is not None:
                        parent.remove(text_node)
                        removed += 1
        return removed

    @staticmethod
    def _isolate_front_matter_pages(document, content_start: int) -> int:
        """Turn front-matter page starts into next-page section boundaries."""
        paragraphs = document.paragraphs
        added = 0
        for paragraph in paragraphs[: max(0, content_start - 1)]:
            sections = paragraph._p.xpath("./w:pPr/w:sectPr")
            if not sections:
                continue
            section_type = sections[0].find(qn("w:type"))
            if section_type is None:
                section_type = OxmlElement("w:type")
                sections[0].insert_element_before(section_type, "w:pgSz")
            section_type.set(qn("w:val"), "nextPage")
        for index in range(1, min(content_start - 1, len(paragraphs))):
            paragraph = paragraphs[index]
            p_pr = paragraph._p.pPr
            if p_pr is None or p_pr.find(qn("w:pageBreakBefore")) is None:
                continue
            previous = paragraphs[index - 1]
            previous_p_pr = previous._p.get_or_add_pPr()
            if previous_p_pr.find(qn("w:sectPr")) is not None:
                p_pr.remove(p_pr.find(qn("w:pageBreakBefore")))
                continue
            following_section = None
            for candidate in paragraphs[index - 1 :]:
                candidates = candidate._p.xpath("./w:pPr/w:sectPr")
                if candidates:
                    following_section = candidates[0]
                    break
            if following_section is None:
                following_section = document.element.body.sectPr
            section = OxmlElement("w:sectPr")
            section_type = OxmlElement("w:type")
            section_type.set(qn("w:val"), "nextPage")
            section.append(section_type)
            # A minimal section inherits headers/footers through Word's normal
            # linkage and avoids duplicating relationship-bound references.
            for name in (
                "pgSz", "pgMar", "paperSrc", "pgBorders", "pgNumType",
                "cols", "vAlign", "titlePg", "textDirection", "docGrid",
            ):
                setting = following_section.find(qn(f"w:{name}"))
                if setting is not None:
                    section.append(deepcopy(setting))
            previous_p_pr.append(section)
            p_pr.remove(p_pr.find(qn("w:pageBreakBefore")))
            added += 1
        return added

    @staticmethod
    def _looks_like_format_specification(document) -> bool:
        heading_text = "".join(
            paragraph.text.strip() for paragraph in document.paragraphs[:12]
        )
        title_signal = bool(
            re.search(r"(?:撰写|写作|排版|格式).{0,8}(?:规范|要求|说明)", heading_text)
        )
        table_labels = {
            row.cells[0].text.strip().replace(" ", "")
            for table in document.tables[:3]
            for row in table.rows
            if row.cells
        }
        rule_labels = {"页面设置", "目录", "正文", "图", "表", "参考文献"}
        return title_signal and len(table_labels & rule_labels) >= 3

    @staticmethod
    def _repair_missing_numbering_part(document, result: ProcessResult) -> None:
        """Repair DOCX files containing dangling numPr references.

        Some Word/WPS documents contain numbering references but omit
        ``word/numbering.xml``. python-docx cannot create that part itself and
        docxcompose consequently raises a message-less NotImplementedError.
        Since the referenced definitions do not exist, remove those dangling
        references and add an empty, valid numbering part before composition.
        """
        try:
            document.part.part_related_by(RT.NUMBERING)
            return
        except KeyError:
            pass

        removed = 0
        for root in (document.element.body, document.styles.element):
            for num_pr in list(root.xpath(".//w:numPr")):
                parent = num_pr.getparent()
                if parent is not None:
                    parent.remove(num_pr)
                    removed += 1
        Composer(document).numbering_part()
        result.warnings.append(
            f"论文缺少编号定义，已移除 {removed} 个无效编号引用并补建编号部件。"
        )

    @staticmethod
    def _trim_body(document, keep_before: int | None = None, remove_before: int | None = None) -> None:
        body = document.element.body
        paragraphs_seen = 0
        for child in list(body.iterchildren()):
            if child.tag == qn("w:sectPr"):
                continue
            if child.tag == qn("w:p"):
                paragraphs_seen += 1
            remove = (keep_before is not None and paragraphs_seen >= keep_before) or (remove_before is not None and paragraphs_seen < remove_before)
            if remove:
                body.remove(child)

    @staticmethod
    def default_output_path(source: Path) -> Path:
        return DocumentProcessor._unique_path(source.with_name(f"{source.stem}-格式修订版.docx"))

    @staticmethod
    def _validate_source(source: Path) -> None:
        if not source.is_file():
            raise FileNotFoundError(f"找不到 Word 文件：{source}")
        if source.suffix.lower() != ".docx":
            raise ValueError("当前版本仅支持 .docx 文件")

    @staticmethod
    def _unique_path(path: Path) -> Path:
        if not path.exists():
            return path
        for index in range(1, 1000):
            candidate = path.with_name(f"{path.stem}_{index}{path.suffix}")
            if not candidate.exists():
                return candidate
        raise FileExistsError(f"无法生成不重名文件：{path}")

    @staticmethod
    def _main_content_start(document) -> int:
        """返回 1 基正文起始段落，保护封面、声明和目录的既有版式。"""
        paragraphs = document.paragraphs
        if not paragraphs:
            return 1
        toc_indexes = []
        for index, paragraph in enumerate(paragraphs, start=1):
            style = paragraph.style
            identity = f"{style.style_id if style else ''} {style.name if style else ''}"
            if re.search(r"(?:^|\s)TOC\s*\d+|目录\s*\d+", identity, re.IGNORECASE):
                toc_indexes.append(index)
        search_from = max(toc_indexes, default=0) + 1
        # A centered cover title such as “（2027）届毕业论文” can resemble a
        # level-3 numbered heading. Anchor the body at a top-level heading so
        # the complete cover/statement section survives the merge.
        for index in range(search_from, len(paragraphs) + 1):
            paragraph = paragraphs[index - 1]
            text = paragraph.text.strip()
            if not text:
                continue
            if DocumentAnalyzer.recognized_heading_level(paragraph) == 1:
                return index
        # 没有标题结构的普通文档仍应从第一段开始处理。
        return 1

    @classmethod
    def _ensure_toc(
        cls, document, rules: DocumentRules, result: ProcessResult, content_start: int
    ) -> int:
        """Insert a real Word TOC field between copied front matter and body."""
        cls._exclude_front_matter_from_toc(document, content_start)
        if document.element.body.xpath(".//w:instrText[contains(., 'TOC ')]"):
            return 0
        paragraphs = document.paragraphs
        if not paragraphs:
            return 0
        anchor = paragraphs[max(0, min(content_start - 1, len(paragraphs) - 1))]._p

        title = document.add_paragraph("目录")
        cls._format_paragraph(title, rules.toc_title)
        # The TOC is inserted exactly at the front/body boundary. A copied
        # front already has a next-page section break; with no copied front,
        # the TOC is the first page. In either case pageBreakBefore is both
        # redundant and capable of creating a numbered blank page.
        title.paragraph_format.page_break_before = None
        title._p.get_or_add_pPr().find(qn("w:outlineLvl")).set(qn("w:val"), "9")

        toc = document.add_paragraph()
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        begin.set(qn("w:dirty"), "true")
        instruction = OxmlElement("w:instrText")
        instruction.set(qn("xml:space"), "preserve")
        instruction.text = ' TOC \\o "1-3" \\h \\z \\u '
        separate = OxmlElement("w:fldChar")
        separate.set(qn("w:fldCharType"), "separate")
        placeholder = OxmlElement("w:t")
        placeholder.text = "正在生成目录…"
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        for child in (begin, instruction, separate, placeholder, end):
            run = OxmlElement("w:r")
            run.append(child)
            toc._p.append(run)

        page_break = document.add_paragraph()
        page_break.add_run().add_break()
        page_break._p.xpath(".//w:br")[-1].set(qn("w:type"), "page")

        for paragraph in (title, toc, page_break):
            anchor.addprevious(paragraph._p)
        result.records.append(
            ChangeRecord(None, "自动目录", "文档中无目录", "封面后插入 1–3 级 Word 目录并刷新页码", "固定系统规则")
        )
        return 3

    @staticmethod
    def _exclude_front_matter_from_toc(document, content_start: int) -> None:
        """Keep cover/declaration paragraphs out of regenerated Word TOCs."""
        for index, paragraph in enumerate(document.paragraphs, start=1):
            if index >= content_start:
                break
            p_pr = paragraph._p.get_or_add_pPr()
            outline = p_pr.find(qn("w:outlineLvl"))
            if outline is None:
                outline = OxmlElement("w:outlineLvl")
                p_pr.append(outline)
            outline.set(qn("w:val"), "9")

    @staticmethod
    def _exclude_non_content_toc_entries(document, start_index: int) -> None:
        excluded = re.compile(
            r"^\s*(?:封面|目录|附录|诚信声明书?|原创性声明|学位论文.{0,8}声明)\s*$",
            re.I,
        )
        for index, paragraph in enumerate(document.paragraphs, start=1):
            if index < start_index or not excluded.match(paragraph.text.strip()):
                continue
            p_pr = paragraph._p.get_or_add_pPr()
            outline = p_pr.find(qn("w:outlineLvl"))
            if outline is None:
                outline = OxmlElement("w:outlineLvl")
                p_pr.append(outline)
            outline.set(qn("w:val"), "9")

    @staticmethod
    def _apply_page_setup(document, rules: DocumentRules, result: ProcessResult) -> None:
        rule = rules.page_setup
        for section_index, section in enumerate(document.sections, start=1):
            old = (
                f"{section.page_width.mm:.1f}×{section.page_height.mm:.1f} mm，"
                f"边距 {section.top_margin.mm:.1f}/{section.bottom_margin.mm:.1f}/"
                f"{section.left_margin.mm:.1f}/{section.right_margin.mm:.1f} mm"
            )
            section.orientation = WD_ORIENT.PORTRAIT
            section.page_width = Mm(rule.width_mm)
            section.page_height = Mm(rule.height_mm)
            section.top_margin = Mm(rule.margin_top_mm)
            section.bottom_margin = Mm(rule.margin_bottom_mm)
            section.left_margin = Mm(rule.margin_left_mm)
            section.right_margin = Mm(rule.margin_right_mm)
            new = (
                f"{rule.width_mm:g}×{rule.height_mm:g} mm，边距 "
                f"{rule.margin_top_mm:g}/{rule.margin_bottom_mm:g}/"
                f"{rule.margin_left_mm:g}/{rule.margin_right_mm:g} mm"
            )
            result.records.append(ChangeRecord(None, f"第 {section_index} 节页面设置", old, new, "启用页面设置规则"))

    @classmethod
    def _enforce_global_paragraph_policy(cls, document, result: ProcessResult) -> None:
        image_count = 0
        for paragraph in document.paragraphs:
            p_pr = paragraph._p.get_or_add_pPr()
            for name in ("widowControl", "keepNext", "keepLines", "pageBreakBefore"):
                cls._set_on_off_property(p_pr, name, False)
            if paragraph._p.xpath(".//w:drawing | .//w:pict"):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                image_count += 1
        result.records.append(ChangeRecord(None, "全局段落策略", "继承原分页属性", "清除段落级换行分页选项", "固定系统规则"))
        if image_count:
            result.records.append(ChangeRecord(None, "图片段落", f"{image_count} 个", "居中、单倍行距", "固定系统规则"))

    @classmethod
    def _apply_normal_text(
        cls,
        document,
        rule: ParagraphRule,
        result: ProcessResult,
        excluded_elements: set[int] | None = None,
        start_index: int = 1,
    ) -> None:
        excluded_elements = excluded_elements or set()
        for index, paragraph in enumerate(document.paragraphs, start=1):
            if index < start_index:
                continue
            if id(paragraph._p) in excluded_elements:
                continue
            if not DocumentAnalyzer.is_normal_body(paragraph):
                continue
            before = cls._paragraph_summary(paragraph)
            cls._format_paragraph(paragraph, rule)
            after = cls._rule_summary(rule)
            result.records.append(
                ChangeRecord(index, "普通正文", before, after, "段落识别为非空普通正文并启用 normal_text 规则")
            )

    @classmethod
    def _apply_headings(
        cls, document, rules: DocumentRules, result: ProcessResult, start_index: int = 1
    ) -> None:
        heading_rules = {
            1: rules.heading_1,
            2: rules.heading_2,
            3: rules.heading_3,
            4: rules.heading_4,
        }
        for index, paragraph in enumerate(document.paragraphs, start=1):
            if index < start_index:
                continue
            if (
                DocumentAnalyzer.is_figure_caption(paragraph)
                or DocumentAnalyzer.is_table_caption(paragraph)
            ):
                continue
            level = DocumentAnalyzer.recognized_heading_level(paragraph)
            if level is None:
                continue
            rule = heading_rules[level]
            if not rule.enabled:
                continue
            before = cls._paragraph_summary(paragraph)
            cls._format_paragraph(paragraph, rule)
            recognition = (
                f"明确的 Heading {level}/标题 {level} 内置样式"
                if DocumentAnalyzer.heading_level(paragraph) == level
                else "编号结构与加粗等标题特征的保守识别"
            )
            result.records.append(
                ChangeRecord(index, f"{level} 级标题", before, cls._rule_summary(rule),
                             f"段落通过{recognition}识别，并启用 heading_{level} 规则")
            )

    @staticmethod
    def _start_chapters_on_new_pages(document, start_index: int, result: ProcessResult) -> None:
        changed = 0
        chapter_pattern = re.compile(
            r"^\s*(?:第\s*[一二三四五六七八九十百零〇0-9]+\s*章|chapter\s+\d+)\b",
            re.I,
        )
        for index, paragraph in enumerate(document.paragraphs, start=1):
            if index < start_index or not chapter_pattern.match(paragraph.text.strip()):
                continue
            if not paragraph.paragraph_format.page_break_before:
                paragraph.paragraph_format.page_break_before = True
                changed += 1
        if changed:
            result.records.append(
                ChangeRecord(None, "章节分页", "章节可能连续排版", f"{changed} 个章节强制另页开始", "每章末尾分页固定规则")
            )

    @classmethod
    def _apply_toc(cls, document, rules: DocumentRules, result: ProcessResult) -> None:
        rule_map = {1: rules.toc_1, 2: rules.toc_2, 3: rules.toc_3}
        for index, paragraph in enumerate(document.paragraphs, start=1):
            text = paragraph.text.strip()
            style = paragraph.style
            identity = f"{style.style_id if style else ''} {style.name if style else ''}"
            if re.fullmatch(r"(?:目\s*录|contents)", text, re.I):
                cls._format_paragraph(paragraph, rules.toc_title)
                paragraph.paragraph_format.page_break_before = True
                # The TOC title is presentation text, not a chapter entry.
                paragraph._p.get_or_add_pPr().find(qn("w:outlineLvl")).set(qn("w:val"), "9")
                result.records.append(ChangeRecord(index, "目录标题", "原格式", cls._rule_summary(rules.toc_title), "目录标题规则"))
                continue
            match = re.search(r"(?:^|\s)(?:TOC|目录)\s*([1-3])", identity, re.I)
            if match:
                rule = rule_map[int(match.group(1))]
                cls._format_paragraph(paragraph, rule)
                result.records.append(ChangeRecord(index, f"目录 {match.group(1)} 级", "原格式", cls._rule_summary(rule), "目录级别样式规则"))

    @classmethod
    def _apply_figure_captions(
        cls,
        document,
        rule: ParagraphRule,
        result: ProcessResult,
        start_index: int = 1,
    ) -> None:
        for index, paragraph in enumerate(document.paragraphs, start=1):
            if index < start_index:
                continue
            if not DocumentAnalyzer.is_figure_caption(paragraph):
                continue
            before = cls._paragraph_summary(paragraph)
            cls._format_paragraph(paragraph, rule)
            result.records.append(
                ChangeRecord(
                    index,
                    "图名",
                    before,
                    cls._rule_summary(rule),
                    "段落通过题注样式或“图/Figure + 编号 + 名称”结构识别，并启用 figure_caption 规则",
                )
            )

    @classmethod
    def _apply_table_captions(
        cls,
        document,
        rule: ParagraphRule,
        result: ProcessResult,
        start_index: int = 1,
    ) -> None:
        for index, paragraph in enumerate(document.paragraphs, start=1):
            if index < start_index:
                continue
            if not DocumentAnalyzer.is_table_caption(paragraph):
                continue
            before = cls._paragraph_summary(paragraph)
            cls._format_paragraph(paragraph, rule)
            result.records.append(
                ChangeRecord(
                    index,
                    "表名",
                    before,
                    cls._rule_summary(rule),
                    "段落通过题注样式或“表/Table + 编号 + 名称”结构识别，并启用 table_caption 规则",
                )
            )

    @classmethod
    def _apply_references(
        cls,
        reference_paragraphs: list[tuple[int, object]],
        rule: ParagraphRule,
        result: ProcessResult,
    ) -> None:
        for index, paragraph in reference_paragraphs:
            before = cls._paragraph_summary(paragraph)
            cls._format_paragraph(paragraph, rule)
            result.records.append(
                ChangeRecord(
                    index,
                    "参考文献条目",
                    before,
                    cls._rule_summary(rule),
                    "段落位于“参考文献/References”标题之后、下一章节标题之前，并启用 reference 规则",
                )
            )

    @staticmethod
    def _reference_paragraphs(
        document, start_index: int = 1
    ) -> list[tuple[int, object]]:
        """返回参考文献标题之后、下一明确章节标题之前的非空段落。"""
        found_heading = False
        references: list[tuple[int, object]] = []
        for index, paragraph in enumerate(document.paragraphs, start=1):
            if index < start_index:
                continue
            text = paragraph.text.strip()
            if DocumentAnalyzer.is_reference_heading(paragraph):
                found_heading = True
                continue
            if not found_heading or not text:
                continue
            if DocumentAnalyzer.is_special_section_heading(text):
                found_heading = False
                continue
            level = DocumentAnalyzer.recognized_heading_level(paragraph)
            if level is not None and not DocumentAnalyzer.is_reference_entry(paragraph):
                found_heading = False
                continue
            references.append((index, paragraph))

        return references

    @classmethod
    def _apply_tables(
        cls,
        document,
        rule: TableRule,
        result: ProcessResult,
        start_index: int = 1,
    ) -> None:
        top_level_tables = cls._tables_in_scope(document, start_index)
        tables = list(cls._iter_table_tree(top_level_tables))
        preserved_equation_tables = 0
        for table_index, table in enumerate(tables, start=1):
            # Word commonly stores a displayed equation and its right-aligned
            # number in a borderless 1x3 table.  It is a layout container, not
            # a data table.  Applying the school's table rule here would expose
            # the hidden grid and can also change the equation's line height.
            # Preserve only the recognized single-row equation layout; a real
            # data table may legitimately contain OMML in one of its cells.
            if cls._is_equation_layout_table(table):
                preserved_equation_tables += 1
                continue
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for floating_property in ("tblInd", "tblpPr"):
                element = table._tbl.tblPr.find(qn(f"w:{floating_property}"))
                if element is not None:
                    table._tbl.tblPr.remove(element)
            cls._set_table_borders(table, rule)
            cls._set_repeat_header_row(table, rule.repeat_header_row)
            if rule.column_width_mm > 0:
                table.autofit = False
                for column in table.columns:
                    column.width = Mm(rule.column_width_mm)
            cell_count = 0
            paragraph_count = 0
            for row in table.rows:
                if rule.row_height_mm > 0:
                    row.height = Mm(rule.row_height_mm)
                    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
                for cell in row.cells:
                    cell_count += 1
                    cell.vertical_alignment = VERTICAL_ALIGNMENTS[rule.vertical_alignment]
                    if rule.column_width_mm > 0:
                        cell.width = Mm(rule.column_width_mm)
                    for paragraph in cell.paragraphs:
                        paragraph_count += 1
                        cls._format_paragraph(paragraph, rule)
            if table.rows:
                for cell in table.rows[0].cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = rule.header_row_bold
            cls._set_table_cell_vertical_alignment(table, rule.vertical_alignment)
            # python-docx's paragraph/run collections do not expose every item
            # nested in hyperlinks or content controls. Apply the locked
            # properties directly to every paragraph/run owned by this table as
            # a final defense against style inheritance.
            cls._enforce_table_paragraph_properties(table)
            cls._enforce_table_run_properties(table, rule)
            result.records.append(
                ChangeRecord(
                    None,
                    f"第 {table_index} 个表格",
                    f"{len(table.rows)} 行 × {len(table.columns)} 列",
                    cls._table_rule_summary(rule),
                    f"启用 table 规则，处理 {cell_count} 个单元格、{paragraph_count} 个段落",
                )
            )
        if preserved_equation_tables:
            result.warnings.append(
                f"已识别并保留 {preserved_equation_tables} 个公式排版表，未套用普通数据表样式。"
            )

    @staticmethod
    def _is_equation_layout_table(table) -> bool:
        """识别经典的“空白 | 公式 | 编号”三栏排版表。"""

        if not table._tbl.xpath(".//m:oMath | .//m:oMathPara"):
            return False
        rows = DocumentProcessor._owned_table_elements(table, "tr")
        if len(rows) != 1:
            return False
        cells = DocumentProcessor._cells_owned_by_row(table, rows[0])
        if len(cells) != 3:
            return False
        formula_cells = [
            index
            for index, cell in enumerate(cells)
            if cell.xpath(".//m:oMath | .//m:oMathPara")
        ]
        if formula_cells != [1]:
            return False
        visible_text = [
            "".join(
                node.text or ""
                for node in cell.xpath(
                    ".//w:t[not(ancestor::m:oMath) and not(ancestor::m:oMathPara)]"
                )
            ).strip()
            for cell in cells
        ]
        if visible_text[0] or visible_text[1]:
            return False
        number = visible_text[2]
        if not number:
            return True
        return bool(
            re.fullmatch(
                r"(?:公式|式)?\s*[（(]?\s*\d+(?:\s*[.．\-—－]\s*\d+)*\s*[)）]?",
                number,
            )
        )

    @staticmethod
    def _tables_in_scope(document, start_index: int):
        """返回正文起点后的顶层表格，包括内容控件包裹的表格。"""

        body_children = list(document.element.body.iterchildren())
        start_position = 0
        if start_index > 1 and document.paragraphs:
            start_element = document.paragraphs[start_index - 1]._p
            try:
                start_position = next(
                    index
                    for index, child in enumerate(body_children)
                    if child is start_element
                )
            except StopIteration:
                start_position = 0

        tables = []
        for child in body_children[start_position:]:
            for element in child.iter(qn("w:tbl")):
                owner_table = next(
                    (
                        ancestor
                        for ancestor in element.iterancestors()
                        if ancestor.tag == qn("w:tbl")
                    ),
                    None,
                )
                if owner_table is None:
                    tables.append(Table(element, document))
        return tables

    @classmethod
    def _iter_table_tree(cls, top_level_tables):
        """Yield each table once, including content-control nested tables."""

        seen_tables = set()

        def visit(table):
            if table._tbl in seen_tables:
                return
            seen_tables.add(table._tbl)
            yield table
            for element in table._tbl.xpath(".//w:tbl"):
                owner_table = next(
                    (
                        ancestor
                        for ancestor in element.iterancestors()
                        if ancestor.tag == qn("w:tbl")
                    ),
                    None,
                )
                if owner_table is table._tbl:
                    yield from visit(Table(element, table._parent))

        for top_level in top_level_tables:
            yield from visit(top_level)

    @staticmethod
    def _owned_table_elements(table, local_name: str):
        """Return descendants whose nearest table ancestor is ``table``."""

        elements = []
        for element in table._tbl.iterdescendants(qn(f"w:{local_name}")):
            owner_table = next(
                (
                    ancestor
                    for ancestor in element.iterancestors()
                    if ancestor.tag == qn("w:tbl")
                ),
                None,
            )
            if owner_table is table._tbl:
                elements.append(element)
        return elements

    @classmethod
    def _cells_owned_by_row(cls, table, row):
        cells = []
        for cell in cls._owned_table_elements(table, "tc"):
            owner_row = next(
                (
                    ancestor
                    for ancestor in cell.iterancestors()
                    if ancestor.tag == qn("w:tr")
                ),
                None,
            )
            if owner_row is row:
                cells.append(cell)
        return cells

    @staticmethod
    def _request_field_update(document, result: ProcessResult) -> None:
        """让 Word/WPS 打开文件时更新目录页码、交叉引用和页码域。"""
        has_fields = bool(document.element.findall(".//" + qn("w:fldChar")))
        if not has_fields:
            for section in document.sections:
                # 访问 section.header/footer 会为原本没有页眉页脚的文件
                # 隐式创建新部件和 relationship。这里只读取已存在的引用，
                # 确保“检查是否有域”本身不会改变文档包结构。
                references = [
                    *section._sectPr.findall(qn("w:headerReference")),
                    *section._sectPr.findall(qn("w:footerReference")),
                ]
                for reference in references:
                    relationship_id = reference.get(qn("r:id"))
                    part = document.part.related_parts.get(relationship_id)
                    if part is not None and part.element.findall(
                        ".//" + qn("w:fldChar")
                    ):
                        has_fields = True
                        break
                if has_fields:
                    break
        if not has_fields:
            return
        settings = document.settings.element
        update = settings.find(qn("w:updateFields"))
        if update is None:
            update = OxmlElement("w:updateFields")
            settings.append(update)
        update.set(qn("w:val"), "true")
        result.records.append(
            ChangeRecord(
                None,
                "域更新设置",
                "未强制更新",
                "打开时自动更新",
                "正文重新排版可能改变页码，保留原域并请求 Word/WPS 更新目录与页码",
            )
        )

    @staticmethod
    def _apply_page_numbers(document, settings: dict, result: ProcessResult) -> None:
        """把原有旧式页码规范为标准 PAGE 域，并保持编号格式切换。"""
        if not settings.get("normalize_existing", False):
            return
        field_sections: list[int] = []
        for index, section in enumerate(document.sections):
            for reference in section._sectPr.findall(qn("w:footerReference")):
                if reference.get(qn("w:type"), "default") != "default":
                    continue
                relationship_id = reference.get(qn("r:id"))
                part = document.part.related_parts.get(relationship_id)
                if part is None:
                    continue
                instructions = part.element.findall(".//" + qn("w:instrText"))
                if any(re.search(r"\bPAGE\b", item.text or "", re.I) for item in instructions):
                    field_sections.append(index)
                    break
        if not field_sections:
            return

        first_numbered = min(field_sections)
        previous_format: str | None = None
        for index in range(first_numbered, len(document.sections)):
            section = document.sections[index]
            pg_num = section._sectPr.find(qn("w:pgNumType"))
            if pg_num is None and index == first_numbered:
                pg_num = OxmlElement("w:pgNumType")
                pg_num.set(qn("w:fmt"), "decimal")
                pg_num.set(qn("w:start"), "1")
                section._sectPr.append(pg_num)

            if pg_num is not None and pg_num.get(qn("w:fmt")):
                current_format = pg_num.get(qn("w:fmt"))
            elif pg_num is not None and pg_num.get(qn("w:start")) is not None:
                current_format = "decimal"
            else:
                current_format = previous_format or "decimal"

            # 同一种编号格式的后续节必须连续；仅格式发生变化（如 I → 1）
            # 时保留显式 start。
            if (
                pg_num is not None
                and index > first_numbered
                and current_format == previous_format
                and pg_num.get(qn("w:start")) is not None
            ):
                del pg_num.attrib[qn("w:start")]
            previous_format = current_format

            footer = section.footer
            footer.is_linked_to_previous = False
            for child in list(footer._element):
                footer._element.remove(child)
            paragraph = footer.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            run.font.name = str(settings.get("font_name", "Times New Roman"))
            run.font.size = Pt(float(settings.get("font_size_pt", 10.5)))
            r_pr = run._element.get_or_add_rPr()
            fonts = r_pr.get_or_add_rFonts()
            fonts.set(qn("w:eastAsia"), "宋体")
            for kind in ("begin", "separate"):
                marker = OxmlElement("w:fldChar")
                marker.set(qn("w:fldCharType"), kind)
                run._r.append(marker)
                if kind == "begin":
                    instruction = OxmlElement("w:instrText")
                    instruction.set(
                        "{http://www.w3.org/XML/1998/namespace}space", "preserve"
                    )
                    instruction.text = " PAGE  \\* MERGEFORMAT "
                    run._r.append(instruction)
            text = OxmlElement("w:t")
            text.text = "1"
            run._r.append(text)
            end = OxmlElement("w:fldChar")
            end.set(qn("w:fldCharType"), "end")
            run._r.append(end)

        result.records.append(
            ChangeRecord(
                None,
                "页码",
                "旧式/不连续页码",
                f"从第 {first_numbered + 1} 节起使用标准居中 PAGE 域并按格式连续编号",
                "模板包含页码域，规范旧式文本框页码并保留罗马数字到阿拉伯数字的切换",
            )
        )

    @classmethod
    def _set_table_borders(cls, table, rule: TableRule) -> None:
        cls._clear_cell_borders(table)
        tbl_pr = table._tbl.tblPr
        old = tbl_pr.find(qn("w:tblBorders"))
        if old is not None:
            tbl_pr.remove(old)
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

        if rule.border_style == "grid":
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                width = rule.outer_border_width_pt if edge in {"top", "left", "bottom", "right"} else rule.inner_border_width_pt
                cls._append_border(borders, edge, "single", width, rule.border_color)
        elif rule.border_style == "three_line":
            cls._append_border(borders, "top", "single", rule.outer_border_width_pt, rule.border_color)
            cls._append_border(borders, "bottom", "single", rule.outer_border_width_pt, rule.border_color)
            for edge in ("left", "right", "insideH", "insideV"):
                cls._append_border(borders, edge, "nil", 0, rule.border_color)
            rows = cls._owned_table_elements(table, "tr")
            if len(rows) > 1:
                for cell in cls._cells_owned_by_row(table, rows[0]):
                    cls._set_cell_bottom_border(cell, rule.inner_border_width_pt, rule.border_color)
        else:
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                cls._append_border(borders, edge, "nil", 0, rule.border_color)

    @classmethod
    def _clear_cell_borders(cls, table) -> None:
        """清除会覆盖 tblBorders 的单元格级残留边框。"""
        # python-docx 会为 row.cells 动态创建代理对象，代理对象的 id 可能在
        # 遍历期间被 Python 复用，不能用 id(cell._tc) 去重。直接遍历物理
        # OOXML 单元格，确保每一个 tcBorders 都被删除。
        for cell in cls._owned_table_elements(table, "tc"):
            tc_pr = cell.find(qn("w:tcPr"))
            borders = None if tc_pr is None else tc_pr.find(qn("w:tcBorders"))
            if borders is not None:
                tc_pr.remove(borders)
        # 行级表格属性例外同样会覆盖表级边框，统一移除其中的边框定义。
        for row in cls._owned_table_elements(table, "tr"):
            row_exception = row.find(qn("w:tblPrEx"))
            borders = (
                None
                if row_exception is None
                else row_exception.find(qn("w:tblBorders"))
            )
            if borders is not None:
                row_exception.remove(borders)

    @classmethod
    def _set_repeat_header_row(cls, table, enabled: bool) -> None:
        """规范跨页表格：仅首行作为重复表头。"""
        rows = cls._owned_table_elements(table, "tr")
        for index, row in enumerate(rows):
            tr_pr = row.find(qn("w:trPr"))
            if tr_pr is None:
                tr_pr = OxmlElement("w:trPr")
                row.insert(0, tr_pr)
            for existing in list(tr_pr.findall(qn("w:tblHeader"))):
                tr_pr.remove(existing)
            if enabled and index == 0:
                header = OxmlElement("w:tblHeader")
                header.set(qn("w:val"), "1")
                tr_pr.append(header)

    @staticmethod
    def _append_border(parent, edge: str, value: str, width_pt: float, color: str) -> None:
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), value)
        if value != "nil":
            element.set(qn("w:sz"), str(max(2, round(width_pt * 8))))
            element.set(qn("w:space"), "0")
            element.set(qn("w:color"), color.lstrip("#").upper())
        parent.append(element)

    @classmethod
    def _set_cell_bottom_border(cls, cell, width_pt: float, color: str) -> None:
        physical_cell = getattr(cell, "_tc", cell)
        tc_pr = physical_cell.get_or_add_tcPr()
        borders = tc_pr.find(qn("w:tcBorders"))
        if borders is None:
            borders = OxmlElement("w:tcBorders")
            tc_pr.append(borders)
        old = borders.find(qn("w:bottom"))
        if old is not None:
            borders.remove(old)
        cls._append_border(borders, "bottom", "single", width_pt, color)

    @classmethod
    def _set_table_cell_vertical_alignment(cls, table, alignment: str) -> None:
        for cell in cls._owned_table_elements(table, "tc"):
            tc_pr = cell.get_or_add_tcPr()
            vertical = tc_pr.get_or_add_vAlign()
            vertical.set(qn("w:val"), alignment)

    @staticmethod
    def _enforce_table_paragraph_properties(table) -> None:
        for paragraph in table._tbl.xpath(".//w:p"):
            owner_table = next(
                (
                    ancestor
                    for ancestor in paragraph.iterancestors()
                    if ancestor.tag == qn("w:tbl")
                ),
                None,
            )
            if owner_table is not table._tbl:
                continue
            p_pr = paragraph.find(qn("w:pPr"))
            if p_pr is None:
                p_pr = OxmlElement("w:pPr")
                paragraph.insert(0, p_pr)
            p_pr.get_or_add_jc().set(qn("w:val"), "center")
            indent = p_pr.get_or_add_ind()
            for name in (
                "left",
                "right",
                "leftChars",
                "rightChars",
                "firstLine",
                "firstLineChars",
                "hanging",
                "hangingChars",
            ):
                indent.set(qn(f"w:{name}"), "0")

    @staticmethod
    def _enforce_table_run_properties(table, rule: TableRule) -> None:
        half_points = str(round(rule.font_size_pt * 2))
        for run in table._tbl.xpath(".//w:r"):
            owner_table = next(
                (
                    ancestor
                    for ancestor in run.iterancestors()
                    if ancestor.tag == qn("w:tbl")
                ),
                None,
            )
            if owner_table is not table._tbl:
                continue
            r_pr = run.find(qn("w:rPr"))
            if r_pr is None:
                r_pr = OxmlElement("w:rPr")
                run.insert(0, r_pr)

            fonts = r_pr.get_or_add_rFonts()
            fonts.set(qn("w:eastAsia"), rule.chinese_font)
            fonts.set(qn("w:ascii"), rule.latin_font)
            fonts.set(qn("w:hAnsi"), rule.latin_font)
            fonts.set(qn("w:cs"), rule.latin_font)
            for theme_attribute in (
                "asciiTheme",
                "hAnsiTheme",
                "eastAsiaTheme",
                "cstheme",
            ):
                qualified = qn(f"w:{theme_attribute}")
                if qualified in fonts.attrib:
                    del fonts.attrib[qualified]

            size = r_pr.get_or_add_sz()
            size.set(qn("w:val"), half_points)
            complex_size = r_pr.find(qn("w:szCs"))
            if complex_size is None:
                complex_size = OxmlElement("w:szCs")
                r_pr.insert(r_pr.index(size) + 1, complex_size)
            complex_size.set(qn("w:val"), half_points)

            for getter in (r_pr.get_or_add_b, r_pr.get_or_add_bCs):
                bold = getter()
                bold.set(qn("w:val"), "0")

    @staticmethod
    def _format_paragraph(paragraph, rule: ParagraphRule) -> None:
        paragraph.alignment = ALIGNMENTS[rule.alignment]
        fmt = paragraph.paragraph_format
        if rule.line_spacing_mode == "fixed":
            fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            fmt.line_spacing = Pt(rule.fixed_line_spacing_pt)
        elif rule.line_spacing_mode == "at_least":
            fmt.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
            fmt.line_spacing = Pt(rule.minimum_line_spacing_pt)
        elif rule.line_spacing_mode == "multiple":
            fmt.line_spacing = rule.multiple_line_spacing
        else:
            fmt.line_spacing_rule = {
                "single": WD_LINE_SPACING.SINGLE,
                "1.5": WD_LINE_SPACING.ONE_POINT_FIVE,
                "double": WD_LINE_SPACING.DOUBLE,
            }[rule.line_spacing_mode]
        DocumentProcessor._set_wps_paragraph_units(paragraph, rule)
        for run in paragraph.runs:
            run.font.name = rule.latin_font
            run.font.size = Pt(rule.font_size_pt)
            run.font.bold = rule.bold
            run.font.italic = rule.italic
            run.font.underline = rule.underline
            r_pr = run._element.get_or_add_rPr()
            fonts = r_pr.get_or_add_rFonts()
            fonts.set(qn("w:eastAsia"), rule.chinese_font)
            fonts.set(qn("w:ascii"), rule.latin_font)
            fonts.set(qn("w:hAnsi"), rule.latin_font)
            spacing = r_pr.find(qn("w:spacing"))
            if rule.character_spacing_mode == "standard" or rule.character_spacing_pt == 0:
                if spacing is not None:
                    r_pr.remove(spacing)
            else:
                if spacing is None:
                    spacing = OxmlElement("w:spacing")
                    r_pr.append(spacing)
                sign = 1 if rule.character_spacing_mode == "expanded" else -1
                spacing.set(qn("w:val"), str(round(sign * rule.character_spacing_pt * 20)))

    @staticmethod
    def _set_wps_paragraph_units(paragraph, rule: ParagraphRule) -> None:
        """直接写入 WPS 会显示为“字符”和“行”的 OOXML 属性。"""
        p_pr = paragraph._p.get_or_add_pPr()
        ind = p_pr.find(qn("w:ind"))
        if ind is None:
            ind = OxmlElement("w:ind")
            p_pr.append(ind)
        for attr in ("firstLine", "hanging", "left", "right", "firstLineChars", "hangingChars", "leftChars", "rightChars"):
            qualified = qn(f"w:{attr}")
            if qualified in ind.attrib:
                del ind.attrib[qualified]
        ind.set(qn("w:left"), str(Cm(rule.left_indent_cm).twips))
        ind.set(qn("w:right"), str(Cm(rule.right_indent_cm).twips))
        if rule.special_indent_mode == "first_line":
            ind.set(qn("w:firstLineChars"), str(round(rule.special_indent_chars * 100)))
        elif rule.special_indent_mode == "hanging":
            ind.set(qn("w:hangingChars"), str(round(rule.special_indent_chars * 100)))
        else:
            # Explicit zeroes are required. Merely omitting these properties
            # allows a paragraph to inherit a first-line or hanging indent from
            # Normal or a custom table style.
            ind.set(qn("w:firstLine"), "0")
            ind.set(qn("w:firstLineChars"), "0")
            ind.set(qn("w:hanging"), "0")
            ind.set(qn("w:hangingChars"), "0")

        spacing = p_pr.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            p_pr.append(spacing)
        for attr in ("before", "after", "beforeLines", "afterLines", "beforeAutospacing", "afterAutospacing"):
            qualified = qn(f"w:{attr}")
            if qualified in spacing.attrib:
                del spacing.attrib[qualified]
        if rule.space_before_unit == "line":
            spacing.set(qn("w:beforeLines"), str(round(rule.space_before_lines * 100)))
        else:
            spacing.set(qn("w:before"), str(round(rule.space_before_pt * 20)))
        if rule.space_after_unit == "line":
            spacing.set(qn("w:afterLines"), str(round(rule.space_after_lines * 100)))
        else:
            spacing.set(qn("w:after"), str(round(rule.space_after_pt * 20)))

        snap = p_pr.find(qn("w:snapToGrid"))
        if snap is None:
            snap = OxmlElement("w:snapToGrid")
            p_pr.append(snap)
        snap.set(qn("w:val"), "1" if rule.snap_to_grid else "0")
        DocumentProcessor._set_on_off_property(p_pr, "adjustRightInd", rule.auto_adjust_right_indent)
        DocumentProcessor._set_on_off_property(p_pr, "bidi", rule.direction == "rtl")
        DocumentProcessor._set_on_off_property(p_pr, "widowControl", rule.widow_control)
        DocumentProcessor._set_on_off_property(p_pr, "keepNext", rule.keep_with_next)
        DocumentProcessor._set_on_off_property(p_pr, "keepLines", rule.keep_lines_together)
        DocumentProcessor._set_on_off_property(p_pr, "pageBreakBefore", rule.page_break_before)
        outline = p_pr.find(qn("w:outlineLvl"))
        if outline is None:
            outline = OxmlElement("w:outlineLvl")
            p_pr.append(outline)
        outline.set(qn("w:val"), str(max(0, min(rule.outline_level, 9))))

    @staticmethod
    def _set_on_off_property(p_pr, name: str, enabled: bool) -> None:
        element = p_pr.find(qn(f"w:{name}"))
        if element is None:
            element = OxmlElement(f"w:{name}")
            p_pr.append(element)
        element.set(qn("w:val"), "1" if enabled else "0")

    @staticmethod
    def _paragraph_summary(paragraph) -> str:
        fmt = paragraph.paragraph_format
        first_run = next((run for run in paragraph.runs if run.text), None)
        font = first_run.font if first_run else None
        return (
            f"样式={paragraph.style.name if paragraph.style else '无'}; "
            f"字体={font.name if font else '继承'}; "
            f"字号={font.size.pt if font and font.size else '继承'}; "
            f"对齐={paragraph.alignment}; 首行缩进={fmt.first_line_indent}"
        )

    @staticmethod
    def _rule_summary(rule: ParagraphRule) -> str:
        spacing = {
            "standard": "标准",
            "expanded": f"加宽 {rule.character_spacing_pt:g}磅",
            "condensed": f"紧缩 {rule.character_spacing_pt:g}磅",
        }[rule.character_spacing_mode]
        line = {
            "single": "单倍行距",
            "1.5": "1.5倍行距",
            "double": "2倍行距",
            "fixed": f"固定值 {rule.fixed_line_spacing_pt:g}磅",
            "at_least": f"最小值 {rule.minimum_line_spacing_pt:g}磅",
            "multiple": f"多倍行距 {rule.multiple_line_spacing:g}倍",
        }[rule.line_spacing_mode]
        before = f"{rule.space_before_lines:g}行" if rule.space_before_unit == "line" else f"{rule.space_before_pt:g}磅"
        after = f"{rule.space_after_lines:g}行" if rule.space_after_unit == "line" else f"{rule.space_after_pt:g}磅"
        return (
            f"中文={rule.chinese_font}; 英文/数字={rule.latin_font}; "
            f"字号={rule.font_size_name}（{rule.font_size_pt:g}磅）; 字距={spacing}; 对齐={rule.alignment}; "
            f"首行缩进={rule.first_line_indent_chars:g}字符; 行距={line}; "
            f"段前/段后={before}/{after}"
        )

    @staticmethod
    def _table_rule_summary(rule: TableRule) -> str:
        border = {"three_line": "三线表", "grid": "全框线", "none": "无框线"}[rule.border_style]
        line = {
            "single": "单倍行距",
            "1.5": "1.5倍行距",
            "double": "2倍行距",
            "fixed": f"固定值 {rule.fixed_line_spacing_pt:g}磅",
            "at_least": f"最小值 {rule.minimum_line_spacing_pt:g}磅",
            "multiple": f"多倍行距 {rule.multiple_line_spacing:g}倍",
        }[rule.line_spacing_mode]
        before = f"{rule.space_before_lines:g}行" if rule.space_before_unit == "line" else f"{rule.space_before_pt:g}磅"
        after = f"{rule.space_after_lines:g}行" if rule.space_after_unit == "line" else f"{rule.space_after_pt:g}磅"
        row_height = f"{rule.row_height_mm:g}mm" if rule.row_height_mm > 0 else "保留/自动"
        column_width = f"{rule.column_width_mm:g}mm" if rule.column_width_mm > 0 else "自动"
        return (
            f"{border}; 中文={rule.chinese_font}; 英文/数字={rule.latin_font}; "
            f"字号={rule.font_size_name}（{rule.font_size_pt:g}磅）; 行距={line}; "
            f"段前/段后={before}/{after}; 对齐={rule.alignment}/"
            f"{rule.vertical_alignment}; 行高={row_height}; 列宽={column_width}"
        )
