from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
import re

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


SPECIAL_SECTIONS = {
    "摘要", "关键词", "abstract", "keywords", "目录", "参考文献", "references",
    "致谢", "结论", "结语", "总结", "结论与展望", "总结与展望",
}
NUMBERED_HEADING = re.compile(r"^\s*(?:第[一二三四五六七八九十]+[章节]|\d+(?:\.\d+){0,3})\s*[^，。；！？]{1,60}$")
CHINESE_LEVELS = {"一": 1, "二": 2, "三": 3, "四": 4}
CAPTION_NUMBER = r"[一二三四五六七八九十百零〇\d]+(?:\s*[-－—.．]\s*[一二三四五六七八九十百零〇\d]+)*"
FIGURE_CAPTION = re.compile(
    rf"^\s*(?:图\s*{CAPTION_NUMBER}|Figure\s*\d+(?:\s*[-.．]\s*\d+)*)"
    r"\s*[：:、.．－—-]?\s*\S+",
    re.IGNORECASE,
)
TABLE_CAPTION = re.compile(
    rf"^\s*(?:(?:续\s*)?表\s*{CAPTION_NUMBER}|Table\s*\d+(?:\s*[-.．]\s*\d+)*)"
    r"\s*[：:、.．－—-]?\s*\S+",
    re.IGNORECASE,
)
REFERENCE_ENTRY = re.compile(r"^\s*(?:\[\s*\d+\s*\]|[（(]?\d+[）).、])")
DRAWINGML_BLIP = "{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
VML_IMAGE_DATA = "{urn:schemas-microsoft-com:vml}imagedata"


@dataclass(slots=True)
class DocumentInfo:
    path: Path
    paragraph_count: int
    non_empty_paragraph_count: int
    table_count: int
    image_count: int
    section_count: int
    headings: list[tuple[int, int | None, str]] = field(default_factory=list)
    uncertain_headings: list[tuple[int, str]] = field(default_factory=list)
    figure_captions: list[tuple[int, str]] = field(default_factory=list)
    table_captions: list[tuple[int, str]] = field(default_factory=list)


class DocumentAnalyzer:
    def analyze(self, path: str | Path) -> DocumentInfo:
        source = Path(path)
        if source.suffix.lower() != ".docx":
            raise ValueError("当前版本仅支持 .docx 文件")
        document = Document(source)
        headings: list[tuple[int, int | None, str]] = []
        uncertain: list[tuple[int, str]] = []
        figure_captions: list[tuple[int, str]] = []
        table_captions: list[tuple[int, str]] = []
        for index, paragraph in enumerate(document.paragraphs, start=1):
            text = paragraph.text.strip()
            if not text:
                continue
            if self.is_figure_caption(paragraph):
                figure_captions.append((index, text))
                continue
            if self.is_table_caption(paragraph):
                table_captions.append((index, text))
                continue
            level = self.recognized_heading_level(paragraph)
            if level is not None:
                headings.append((index, level, text))
            elif self.is_probable_heading(paragraph):
                uncertain.append((index, text))
        # 兼容现代 DrawingML 图片和旧论文中常见的 VML `w:pict` 图片。
        # `document.inline_shapes` 只覆盖前者，会把真实旧稿误报为 0 张图。
        image_count = len(document.element.findall(f".//{DRAWINGML_BLIP}")) + len(
            document.element.findall(f".//{VML_IMAGE_DATA}")
        )
        return DocumentInfo(
            path=source,
            paragraph_count=len(document.paragraphs),
            non_empty_paragraph_count=sum(bool(p.text.strip()) for p in document.paragraphs),
            table_count=len(document.tables),
            image_count=image_count,
            section_count=len(document.sections),
            headings=headings,
            uncertain_headings=uncertain,
            figure_captions=figure_captions,
            table_captions=table_captions,
        )

    @classmethod
    def heading_level(cls, paragraph: Paragraph) -> int | None:
        """返回明确的大纲级别，兼容内置、自定义样式和直接 outlineLvl。"""
        direct = cls._outline_level(paragraph._p.pPr)
        if direct is not None:
            return direct
        style = paragraph.style
        style_id = (style.style_id if style else "") or ""
        name = (style.name if style else "") or ""
        identity = f"{style_id} {name}"
        match = re.search(r"(?:Heading|标题)\s*([1-4])(?!\d)", identity, re.IGNORECASE)
        if match:
            return int(match.group(1))
        match = re.search(r"([一二三四])\s*级\s*标题|标题\s*([一二三四])\s*级", identity)
        if match:
            return CHINESE_LEVELS[match.group(1) or match.group(2)]
        if re.search(r"小节\s*标题", identity):
            return 3
        if re.search(r"(?:章节|章)\s*标题", identity):
            return 1
        if re.search(r"(?<!小)节\s*标题", identity):
            return 2
        current = style
        visited = set()
        while current is not None and id(current) not in visited:
            visited.add(id(current))
            level = cls._outline_level(current._element.pPr)
            if level is not None:
                return level
            current = current.base_style
        return None

    @staticmethod
    def _outline_level(p_pr) -> int | None:
        if p_pr is None:
            return None
        outline = p_pr.find(qn("w:outlineLvl"))
        if outline is None:
            return None
        try:
            value = int(outline.get(qn("w:val")))
        except (TypeError, ValueError):
            return None
        return value + 1 if 0 <= value <= 3 else None

    @classmethod
    def recognized_heading_level(cls, paragraph: Paragraph) -> int | None:
        """先用大纲/样式，再对具有标题外观的编号段落保守推断。"""
        if cls.caption_kind(paragraph) is not None:
            return None
        styled = cls.heading_level(paragraph)
        if styled is not None:
            return styled
        text = paragraph.text.strip()
        if not text or len(text) > 80:
            return None
        if cls.is_special_section_heading(text):
            return 1
        if re.match(r"^第[一二三四五六七八九十百\d]+章(?:\s|$)", text):
            return 1
        if re.match(r"^第[一二三四五六七八九十百\d]+节(?:\s|$)", text):
            return 2 if cls._has_heading_appearance(paragraph) else None
        if not cls._has_heading_appearance(paragraph):
            return None
        # “一、/（一）/1.”是中文论文常见的三级层次；单独的“1.”不应
        # 与阿拉伯章节制中的“1 标题”混为一级标题。
        if re.match(r"^\d+\s*[.．]\s*(?!\d)\S", text):
            return 3
        decimal = re.match(r"^(\d+(?:[.．]\d+){0,3})(?:\s+|[、．.)）])\s*\S", text)
        if decimal:
            return min(len(re.split(r"[.．]", decimal.group(1))), 4)
        if re.match(r"^[一二三四五六七八九十百]+、\s*\S", text):
            return 1
        if re.match(r"^[（(][一二三四五六七八九十百]+[）)]\s*\S", text):
            return 2
        if re.match(r"^[（(]\d+[）)]\s*\S", text):
            return 3
        return None

    @classmethod
    def _has_heading_appearance(cls, paragraph: Paragraph) -> bool:
        """要求至少一个视觉/分页标题信号，避免把普通编号列表当作标题。"""
        if cls._is_list_paragraph(paragraph):
            return False
        text = paragraph.text.strip()
        if re.search(r"[。！？；;]$", text):
            return False
        visible_runs = [run for run in paragraph.runs if run.text.strip()]
        if visible_runs and sum(run.bold is True for run in visible_runs) >= len(visible_runs) / 2:
            return True
        explicit_sizes = [run.font.size.pt for run in visible_runs if run.font.size is not None]
        if explicit_sizes and max(explicit_sizes) >= 14:
            return True
        style = paragraph.style
        if style is not None:
            if style.font.bold is True or (style.font.size is not None and style.font.size.pt >= 14):
                return True
        fmt = paragraph.paragraph_format
        style_fmt = style.paragraph_format if style is not None else None
        return bool(
            paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
            or (style_fmt is not None and style_fmt.alignment == WD_ALIGN_PARAGRAPH.CENTER)
            or fmt.keep_with_next is True
            or fmt.page_break_before is True
            or (style_fmt is not None and style_fmt.keep_with_next is True)
            or (style_fmt is not None and style_fmt.page_break_before is True)
        )

    @staticmethod
    def _is_list_paragraph(paragraph: Paragraph) -> bool:
        style = paragraph.style
        identity = f"{style.style_id if style else ''} {style.name if style else ''}"
        if re.search(r"List|列表|项目符号|项目编号", identity, re.IGNORECASE):
            return True
        p_pr = paragraph._p.pPr
        return p_pr is not None and p_pr.find(qn("w:numPr")) is not None

    @classmethod
    def is_probable_heading(cls, paragraph: Paragraph) -> bool:
        text = paragraph.text.strip()
        if not text or len(text) > 80 or not NUMBERED_HEADING.match(text):
            return False
        return cls._has_heading_appearance(paragraph)

    @classmethod
    def caption_kind(cls, paragraph: Paragraph) -> str | None:
        text = paragraph.text.strip()
        if not text or len(text) > 200:
            return None
        if FIGURE_CAPTION.match(text):
            return "figure"
        if TABLE_CAPTION.match(text):
            return "table"
        # 个别原稿题注会漏写“图”字；仅在前一段确实含图片、当前段居中且
        # 文本是“章号.图号 名称”时容错，避免把普通二级标题误判成图题。
        previous = paragraph._p.getprevious()
        style = paragraph.style
        style_alignment = style.paragraph_format.alignment if style is not None else None
        if (
            previous is not None
            and (
                previous.find(".//" + qn("w:drawing")) is not None
                or previous.find(".//" + qn("w:pict")) is not None
            )
            and re.match(r"^\s*\d+\s*[.．-]\s*\d+\s+\S+", text)
            and (
                paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
                or style_alignment == WD_ALIGN_PARAGRAPH.CENTER
            )
        ):
            return "figure"
        style_text = f"{style.style_id if style else ''} {style.name if style else ''}"
        if re.search(r"Caption|题注", style_text, re.IGNORECASE):
            if re.match(r"^\s*(?:图(?!中|如|所示)|Figure\b)", text, re.IGNORECASE):
                return "figure"
            if re.match(r"^\s*(?:表(?!中|格)|Table\b)", text, re.IGNORECASE):
                return "table"
        return None

    @classmethod
    def is_figure_caption(cls, paragraph: Paragraph) -> bool:
        return cls.caption_kind(paragraph) == "figure"

    @classmethod
    def is_table_caption(cls, paragraph: Paragraph) -> bool:
        return cls.caption_kind(paragraph) == "table"

    @staticmethod
    def is_reference_heading(paragraph_or_text: Paragraph | str) -> bool:
        text = paragraph_or_text.text if isinstance(paragraph_or_text, Paragraph) else paragraph_or_text
        normalized = re.sub(r"\s+", "", text).strip("[]【】()（）:：")
        return normalized.casefold() in {"参考文献", "references"}

    @staticmethod
    def is_reference_entry(paragraph: Paragraph) -> bool:
        return bool(REFERENCE_ENTRY.match(paragraph.text))

    @staticmethod
    def is_special_section_heading(text: str) -> bool:
        normalized = re.sub(r"\s+", "", text).strip("[]【】()（）:：")
        folded = normalized.casefold()
        if folded in SPECIAL_SECTIONS:
            return True
        return bool(
            re.fullmatch(r"(?:结论|总结)(?:与展望)?", normalized)
            or re.fullmatch(r"附录(?:[A-Za-z一二三四五六七八九十\d]+)?", normalized)
        )

    @classmethod
    def is_normal_body(cls, paragraph: Paragraph) -> bool:
        text = paragraph.text.strip()
        if not text:
            return False
        if cls.is_figure_caption(paragraph):
            return False
        if cls.is_table_caption(paragraph):
            return False
        if cls.recognized_heading_level(paragraph) is not None:
            return False
        return not cls.is_probable_heading(paragraph)
