from __future__ import annotations

from pathlib import Path
import sys
import tempfile
import unittest

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Pt
from docx.table import Table


TOOL_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(TOOL_ROOT))

from word_formatter.core.processor import DocumentProcessor  # noqa: E402
from word_formatter.models.results import ProcessResult  # noqa: E402
from word_formatter.models.rules import DocumentRules  # noqa: E402


class DocumentProcessorCompatibilityTests(unittest.TestCase):
    def test_repairs_dangling_numbering_references_before_composition(self) -> None:
        document = Document()
        paragraph = document.add_paragraph("带有损坏编号引用的正文")
        num_pr = OxmlElement("w:numPr")
        num_id = OxmlElement("w:numId")
        num_id.set(qn("w:val"), "9")
        num_pr.append(num_id)
        paragraph._p.get_or_add_pPr().append(num_pr)
        numbering_rel = next(
            rel for rel in document.part.rels.values() if rel.reltype == RT.NUMBERING
        )
        del document.part.rels[numbering_rel.rId]
        result = ProcessResult(Path("source.docx"), Path("output.docx"))

        DocumentProcessor._repair_missing_numbering_part(document, result)

        self.assertFalse(document.element.body.xpath(".//w:numPr"))
        self.assertIsNotNone(document.part.part_related_by(RT.NUMBERING))
        self.assertTrue(any("缺少编号定义" in item for item in result.warnings))


class DocumentProcessorTableTests(unittest.TestCase):
    def assert_locked_table_core(self, table) -> None:
        self.assertEqual(table.alignment, WD_TABLE_ALIGNMENT.CENTER)
        borders = table._tbl.tblPr.find(qn("w:tblBorders"))
        self.assertIsNotNone(borders)
        assert borders is not None
        self.assertEqual(borders.find(qn("w:top")).get(qn("w:val")), "single")
        self.assertEqual(borders.find(qn("w:top")).get(qn("w:sz")), "12")
        self.assertEqual(
            borders.find(qn("w:bottom")).get(qn("w:val")), "single"
        )
        self.assertEqual(borders.find(qn("w:bottom")).get(qn("w:sz")), "12")
        for edge in ("left", "right", "insideH", "insideV"):
            self.assertEqual(borders.find(qn(f"w:{edge}")).get(qn("w:val")), "nil")
        for row in table.rows:
            for cell in row.cells:
                self.assertEqual(
                    cell.vertical_alignment, WD_CELL_VERTICAL_ALIGNMENT.CENTER
                )
                for paragraph in cell.paragraphs:
                    self.assertEqual(paragraph.alignment, WD_ALIGN_PARAGRAPH.CENTER)
                    ind = paragraph._p.get_or_add_pPr().find(qn("w:ind"))
                    self.assertIsNotNone(ind)
                    assert ind is not None
                    self.assertEqual(ind.get(qn("w:firstLine")), "0")
                    self.assertEqual(ind.get(qn("w:firstLineChars")), "0")

    def test_equation_layout_table_is_preserved_while_data_table_is_formatted(self) -> None:
        with tempfile.TemporaryDirectory(prefix="dokiai_processor_") as directory:
            root = Path(directory)
            source = root / "source.docx"
            output = root / "output.docx"

            document = Document()
            document.styles["Normal"].paragraph_format.first_line_indent = Pt(52.2)
            equation_table = document.add_table(rows=1, cols=3)
            equation_paragraph = equation_table.cell(0, 1).paragraphs[0]
            equation_paragraph._p.append(OxmlElement("m:oMath"))
            equation_table.cell(0, 2).text = "(1-1)"

            data_table = document.add_table(rows=2, cols=2)
            data_table.style = "Table Grid"
            floating = OxmlElement("w:tblpPr")
            floating.set(qn("w:tblpX"), "720")
            data_table._tbl.tblPr.append(floating)
            data_table.cell(0, 0).text = "名称"
            data_table.cell(0, 1).text = "数值"
            data_table.cell(1, 0).text = "速度"
            data_table.cell(1, 1).text = "10"
            for row in data_table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        paragraph.paragraph_format.first_line_indent = Pt(24)
                        for run in paragraph.runs:
                            run.font.name = "Arial"
                            run.font.size = Pt(22)
                            run.font.bold = True
                            fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
                            for theme_slot in (
                                "asciiTheme",
                                "hAnsiTheme",
                                "eastAsiaTheme",
                                "cstheme",
                            ):
                                fonts.set(qn(f"w:{theme_slot}"), "majorHAnsi")
            document.save(source)

            rules = DocumentRules()
            rules.page_setup.enabled = False
            rules.normal_text.enabled = False
            rules.table.enabled = False
            rules.table.chinese_font = "仿宋"
            rules.table.latin_font = "Arial"
            rules.table.number_font = "Arial"
            rules.table.font_size_name = "二号"
            rules.table.font_size_pt = 22.0
            rules.table.bold = True
            rules.table.header_row_bold = True
            rules.table.alignment = "justify"
            rules.table.left_indent_cm = 1.0
            rules.table.right_indent_cm = 1.0
            rules.table.special_indent_mode = "first_line"
            rules.table.special_indent_chars = 2.0
            rules.table.border_style = "grid"
            rules.table.border_color = "FF0000"
            rules.table.outer_border_width_pt = 0.25
            rules.table.inner_border_width_pt = 0.25
            rules.table.vertical_alignment = "top"
            rules.table.repeat_header_row = True

            result = DocumentProcessor().process(source, rules, output)

            formatted = Document(output)
            formatted_equation = formatted.tables[0]
            formatted_data = formatted.tables[1]

            self.assertTrue(
                DocumentProcessor._is_equation_layout_table(formatted_equation)
            )
            self.assertIsNone(
                formatted_equation._tbl.tblPr.find(qn("w:tblBorders"))
            )
            self.assertFalse(
                formatted_equation._tbl.xpath(".//w:trPr/w:tblHeader")
            )
            self.assertEqual(formatted_data.alignment, WD_TABLE_ALIGNMENT.CENTER)
            self.assertIsNone(formatted_data._tbl.tblPr.find(qn("w:tblInd")))
            self.assertIsNone(formatted_data._tbl.tblPr.find(qn("w:tblpPr")))

            borders = formatted_data._tbl.tblPr.find(qn("w:tblBorders"))
            self.assertIsNotNone(borders)
            assert borders is not None
            expected_edges = {
                "top": ("single", "12"),
                "bottom": ("single", "12"),
                "left": ("nil", None),
                "right": ("nil", None),
                "insideH": ("nil", None),
                "insideV": ("nil", None),
            }
            for edge, (value, size) in expected_edges.items():
                element = borders.find(qn(f"w:{edge}"))
                self.assertIsNotNone(element, edge)
                assert element is not None
                self.assertEqual(element.get(qn("w:val")), value, edge)
                self.assertEqual(element.get(qn("w:sz")), size, edge)

            for cell in formatted_data.rows[0].cells:
                bottom = cell._tc.xpath("./w:tcPr/w:tcBorders/w:bottom")
                self.assertEqual(len(bottom), 1)
                self.assertEqual(bottom[0].get(qn("w:val")), "single")
                self.assertEqual(bottom[0].get(qn("w:sz")), "6")

            self.assertTrue(formatted_data._tbl.xpath(".//w:trPr/w:tblHeader"))
            for row in formatted_data.rows:
                for cell in row.cells:
                    self.assertEqual(
                        cell.vertical_alignment, WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    )
                    for paragraph in cell.paragraphs:
                        self.assertEqual(
                            paragraph.alignment, WD_ALIGN_PARAGRAPH.CENTER
                        )
                        ind = paragraph._p.get_or_add_pPr().find(qn("w:ind"))
                        self.assertIsNotNone(ind)
                        assert ind is not None
                        for name in (
                            "left",
                            "right",
                            "firstLine",
                            "firstLineChars",
                            "hanging",
                            "hangingChars",
                        ):
                            self.assertEqual(ind.get(qn(f"w:{name}")), "0", name)

            for run in formatted_data._tbl.xpath(".//w:r"):
                r_pr = run.find(qn("w:rPr"))
                self.assertIsNotNone(r_pr)
                assert r_pr is not None
                fonts = r_pr.find(qn("w:rFonts"))
                self.assertIsNotNone(fonts)
                assert fonts is not None
                for slot in ("eastAsia", "ascii", "hAnsi", "cs"):
                    expected_font = "宋体" if slot == "eastAsia" else "Times New Roman"
                    self.assertEqual(fonts.get(qn(f"w:{slot}")), expected_font, slot)
                for theme_slot in (
                    "asciiTheme",
                    "hAnsiTheme",
                    "eastAsiaTheme",
                    "cstheme",
                ):
                    self.assertIsNone(fonts.get(qn(f"w:{theme_slot}")))
                self.assertEqual(r_pr.find(qn("w:sz")).get(qn("w:val")), "24")
                self.assertEqual(r_pr.find(qn("w:szCs")).get(qn("w:val")), "24")
                self.assertEqual(r_pr.find(qn("w:b")).get(qn("w:val")), "0")
                self.assertEqual(r_pr.find(qn("w:bCs")).get(qn("w:val")), "0")

            self.assertTrue(rules.table.enabled)
            self.assertEqual(rules.table.border_style, "three_line")
            self.assertEqual(rules.table.chinese_font, "宋体")
            self.assertEqual(rules.table.latin_font, "Times New Roman")
            self.assertEqual(rules.table.font_size_pt, 12.0)
            self.assertFalse(rules.table.bold)
            self.assertFalse(rules.table.header_row_bold)
            self.assertEqual(rules.table.alignment, "center")
            self.assertEqual(rules.table.special_indent_mode, "none")
            self.assertIn(
                "已识别并保留 1 个公式排版表，未套用普通数据表样式。",
                result.warnings,
            )

    def test_formula_data_nested_and_content_control_tables_are_locked(self) -> None:
        with tempfile.TemporaryDirectory(prefix="dokiai_processor_") as directory:
            root = Path(directory)
            source = root / "source.docx"
            output = root / "output.docx"

            document = Document()
            formula_data = document.add_table(rows=2, cols=2)
            formula_data.cell(0, 0).text = "项目"
            formula_data.cell(0, 1).text = "计算式"
            formula_data.cell(1, 0).text = "速度"
            protected_paragraph = formula_data.cell(1, 0).paragraphs[0]
            protected_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            protected_paragraph.paragraph_format.first_line_indent = Pt(24)
            protected_paragraph.runs[0].font.bold = True
            formula_data.cell(1, 0).add_paragraph()
            protected_element = protected_paragraph._p
            protected_parent = protected_element.getparent()
            protected_position = protected_parent.index(protected_element)
            protected_parent.remove(protected_element)
            protected_sdt = OxmlElement("w:sdt")
            protected_content = OxmlElement("w:sdtContent")
            protected_content.append(protected_element)
            protected_sdt.append(protected_content)
            protected_parent.insert(protected_position, protected_sdt)
            formula_data.cell(1, 1).paragraphs[0]._p.append(OxmlElement("m:oMath"))

            outer = document.add_table(rows=2, cols=1)
            outer.cell(0, 0).text = "外层表头"
            nested = outer.cell(1, 0).add_table(rows=2, cols=1)
            nested.style = "Table Grid"
            nested.cell(0, 0).text = "嵌套表头"
            nested.cell(1, 0).text = "嵌套内容"

            single_row = document.add_table(rows=1, cols=2)
            single_row.cell(0, 0).paragraphs[0]._p.append(OxmlElement("m:oMath"))
            single_row.cell(0, 1).text = "5"

            wrapped_row_table = document.add_table(rows=2, cols=1)
            wrapped_row_table.cell(0, 0).text = "行内容控件表头"
            wrapped_row_table.cell(1, 0).text = "行内容控件内容"
            wrapped_row_element = wrapped_row_table.rows[0]._tr
            wrapped_row_parent = wrapped_row_element.getparent()
            wrapped_row_position = wrapped_row_parent.index(wrapped_row_element)
            wrapped_row_parent.remove(wrapped_row_element)
            row_sdt = OxmlElement("w:sdt")
            row_sdt_content = OxmlElement("w:sdtContent")
            row_sdt_content.append(wrapped_row_element)
            row_sdt.append(row_sdt_content)
            wrapped_row_parent.insert(wrapped_row_position, row_sdt)

            wrapped = document.add_table(rows=2, cols=1)
            wrapped.cell(0, 0).text = "内容控件表头"
            wrapped.cell(1, 0).text = "内容控件表格"
            wrapped_element = wrapped._tbl
            body = wrapped_element.getparent()
            wrapped_position = body.index(wrapped_element)
            body.remove(wrapped_element)
            sdt = OxmlElement("w:sdt")
            sdt_content = OxmlElement("w:sdtContent")
            sdt_content.append(wrapped_element)
            sdt.append(sdt_content)
            body.insert(wrapped_position, sdt)
            document.save(source)

            rules = DocumentRules()
            rules.page_setup.enabled = False
            rules.normal_text.enabled = False
            rules.table.border_style = "grid"
            result = DocumentProcessor().process(source, rules, output)

            formatted = Document(output)
            formatted_formula = formatted.tables[0]
            formatted_outer = formatted.tables[1]
            formatted_nested = formatted_outer.cell(1, 0).tables[0]
            formatted_single = formatted.tables[2]
            formatted_wrapped_row = formatted.tables[3]
            wrapped_tables = formatted.element.body.xpath(".//w:sdtContent/w:tbl")
            self.assertEqual(len(wrapped_tables), 1)
            formatted_wrapped = Table(wrapped_tables[0], formatted)

            self.assertFalse(
                DocumentProcessor._is_equation_layout_table(formatted_formula)
            )
            self.assertFalse(
                DocumentProcessor._is_equation_layout_table(formatted_single)
            )
            for table in (
                formatted_formula,
                formatted_outer,
                formatted_nested,
                formatted_single,
                formatted_wrapped_row,
                formatted_wrapped,
            ):
                self.assert_locked_table_core(table)

            self.assertFalse(
                formatted_single._tbl.xpath(
                    "./w:tr[1]/w:tc/w:tcPr/w:tcBorders/w:bottom"
                )
            )
            wrapped_rows = formatted_wrapped_row._tbl.xpath(
                ".//w:sdtContent/w:tr"
            )
            self.assertEqual(len(wrapped_rows), 1)
            self.assertTrue(wrapped_rows[0].xpath("./w:trPr/w:tblHeader"))
            wrapped_cells = wrapped_rows[0].xpath(".//w:tc")
            self.assertEqual(len(wrapped_cells), 1)
            self.assertEqual(
                wrapped_cells[0]
                .find("./" + qn("w:tcPr") + "/" + qn("w:vAlign"))
                .get(qn("w:val")),
                "center",
            )
            wrapped_header_bottom = wrapped_cells[0].xpath(
                "./w:tcPr/w:tcBorders/w:bottom"
            )
            self.assertEqual(len(wrapped_header_bottom), 1)
            self.assertEqual(
                wrapped_header_bottom[0].get(qn("w:sz")), "6"
            )
            protected_paragraphs = formatted_formula._tbl.xpath(
                ".//w:sdtContent/w:p"
            )
            self.assertEqual(len(protected_paragraphs), 1)
            protected_p_pr = protected_paragraphs[0].find(qn("w:pPr"))
            self.assertIsNotNone(protected_p_pr)
            assert protected_p_pr is not None
            self.assertEqual(
                protected_p_pr.find(qn("w:jc")).get(qn("w:val")), "center"
            )
            protected_indent = protected_p_pr.find(qn("w:ind"))
            self.assertIsNotNone(protected_indent)
            assert protected_indent is not None
            self.assertTrue(
                all(
                    protected_indent.get(qn(f"w:{name}")) == "0"
                    for name in (
                        "left",
                        "right",
                        "leftChars",
                        "rightChars",
                        "firstLine",
                        "firstLineChars",
                        "hanging",
                        "hangingChars",
                    )
                )
            )
            protected_r_pr = protected_paragraphs[0].find(".//" + qn("w:rPr"))
            self.assertIsNotNone(protected_r_pr)
            assert protected_r_pr is not None
            self.assertEqual(
                protected_r_pr.find(qn("w:b")).get(qn("w:val")), "0"
            )
            self.assertFalse(
                formatted_nested._tbl.xpath(
                    "./w:tr/w:tc/w:tcPr/w:tcBorders/w:left"
                )
            )
            self.assertFalse(
                any("公式排版表" in warning for warning in result.warnings)
            )


if __name__ == "__main__":
    unittest.main()
