from pathlib import Path
import sys
import tempfile
import unittest
from zipfile import ZipFile

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import RGBColor

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))
from word_formatter.core.template_text import read_template_text


class TemplateTextTests(unittest.TestCase):
    def read(self, document):
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "template.docx"
            document.save(path)
            return read_template_text(path)

    def test_specification_table_is_evidence_not_a_cover(self):
        document = Document()
        document.add_paragraph("附件3：毕业设计（论文）撰写规范")
        document.add_paragraph("（一）封面")
        document.add_paragraph("采用学校规定的统一格式封面（见附件一），按要求填写题目、姓名、学号、专业。")
        document.add_paragraph("（二）诚信承诺书")
        document.add_paragraph("采用学校统一提供的内容（见附件二）。")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "正文"
        table.cell(0, 1).text = "中文宋体小四号；英文 Times New Roman 小四号；1.25倍行距。"
        table.cell(1, 0).text = "一级标题"
        table.cell(1, 1).text = "黑体小二号，段前段后20磅。"
        evidence = self.read(document)
        self.assertEqual(evidence["documentKindHint"], "specification")
        self.assertIsNone(evidence["copyCandidate"])
        table_text = next(block["text"] for block in evidence["textBlocks"] if block["kind"] == "table")
        self.assertIn("段前段后20磅", table_text)
        self.assertIn("Times New Roman", table_text)

    def test_actual_cover_and_statement_prefix_ends_before_abstract(self):
        document = Document()
        document.add_paragraph("某大学本科毕业论文")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "姓名"
        table.cell(0, 1).text = "张三"
        table.cell(1, 0).text = "学号"
        table.cell(1, 1).text = "2026001"
        document.add_paragraph("学术诚信声明")
        document.add_paragraph("本人郑重声明：本论文是在导师的指导下独立进行研究工作并完成的研究成果，引用内容均已注明来源。")
        document.add_paragraph("摘 要")
        document.add_paragraph("这里是摘要内容。")
        evidence = self.read(document)
        self.assertEqual(evidence["documentKindHint"], "template")
        self.assertEqual(evidence["copyCandidate"]["startParagraph"], 1)
        self.assertEqual(evidence["copyCandidate"]["endParagraph"], 3)
        self.assertGreaterEqual(len(evidence["copyCandidate"]["evidenceIds"]), 3)

    def test_spaced_cover_fields_and_standalone_cover_are_recognized(self):
        document = Document()
        document.add_paragraph("本科毕业论文")
        document.add_paragraph("姓    名：张三")
        document.add_paragraph("学    号：2026001")
        evidence = self.read(document)
        self.assertEqual(evidence["documentKindHint"], "template")
        self.assertEqual(evidence["copyCandidate"]["endParagraph"], 3)

    def test_long_document_without_verified_body_boundary_does_not_copy_everything(self):
        document = Document()
        document.add_paragraph("本科毕业论文")
        document.add_paragraph("姓名：张三")
        document.add_paragraph("学号：2026001")
        for _ in range(90):
            document.add_paragraph("研究内容。")
        self.assertIsNone(self.read(document)["copyCandidate"])

    def test_order_merged_cells_red_text_textbox_and_package_annotations(self):
        document = Document()
        first = document.add_paragraph("规则在这里")
        red = first.add_run("红字：正文为宋体小四号")
        red.font.color.rgb = RGBColor(255, 0, 0)
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).merge(table.cell(0, 1)).text = "合并单元格只出现一次"
        table.cell(1, 0).text = "段前"
        table.cell(1, 1).text = "20磅"
        last = document.add_paragraph("最后一个正文段落")
        textbox = OxmlElement("w:txbxContent")
        textbox_p = OxmlElement("w:p")
        textbox_r = OxmlElement("w:r")
        textbox_t = OxmlElement("w:t")
        textbox_t.text = "文本框：三级标题黑体四号"
        textbox_r.append(textbox_t)
        textbox_p.append(textbox_r)
        textbox.append(textbox_p)
        last._p.append(textbox)
        document.sections[0].header.paragraphs[0].text = "学校页眉要求"
        document.sections[0].footer.paragraphs[0].text = "页脚页码居中"
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "comments.docx"
            document.save(path)
            comment_xml = '<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:comment w:id="0"><w:p><w:r><w:t>批注要求：图标题宋体五号</w:t></w:r></w:p></w:comment></w:comments>'
            with ZipFile(path, "a") as package:
                package.writestr("word/comments.xml", comment_xml.encode("utf-8"))
            evidence = read_template_text(path)
        blocks = evidence["textBlocks"]
        self.assertEqual([block["kind"] for block in blocks[:3]], ["paragraph", "table", "paragraph"])
        self.assertEqual([block["paragraphStart"] for block in blocks[:3]], [1, 1, 2])
        text = "\n".join(block["text"] for block in blocks)
        self.assertEqual(text.count("合并单元格只出现一次"), 1)
        self.assertEqual(text.count("文本框：三级标题黑体四号"), 1)
        self.assertIn("红字：正文为宋体小四号", text)
        self.assertIn("批注要求：图标题宋体五号", text)
        self.assertEqual({block["kind"] for block in blocks}, {"paragraph", "table", "comment", "header", "footer"})

    def test_large_sample_prioritizes_late_explicit_rules_and_reports_omissions(self):
        document = Document()
        document.add_paragraph("毕业设计论文撰写规范")
        for _ in range(260):
            document.add_paragraph("这是示例论文内容，仅用于展示研究内容。" * 30)
        document.add_paragraph("三级标题：黑体四号，1.5倍行距，段前20磅。")
        evidence = self.read(document)
        self.assertIn("三级标题：黑体四号", "\n".join(block["text"] for block in evidence["textBlocks"]))
        self.assertLessEqual(sum(len(block["text"]) for block in evidence["textBlocks"]), 48_000)
        self.assertTrue(any("省略" in note for note in evidence["notes"]))


if __name__ == "__main__":
    unittest.main()
