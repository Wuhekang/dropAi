from __future__ import annotations

import hashlib
import json
import os
from pathlib import Path
import subprocess
import sys
import tempfile
import unittest
from argparse import Namespace
from unittest.mock import patch

from docx import Document


TOOL_ROOT = Path(__file__).resolve().parents[1]
CLI = TOOL_ROOT / "format_cli.py"
sys.path.insert(0, str(TOOL_ROOT))

from format_cli import (  # noqa: E402
    _apply_confirmed_rules,
    _editable_rules,
    _publish_without_overwrite,
    run_job,
    validate_runtime_support,
)
import format_cli  # noqa: E402
from word_formatter.core.word_converter import WordConversionError  # noqa: E402
from word_formatter.models.rules import DocumentRules, enforce_locked_document_policy  # noqa: E402


def file_hash(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def visible_text(path: Path) -> list[str]:
    document = Document(path)
    values = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            values.extend(cell.text for cell in row.cells)
    return values


def build_template(path: Path) -> None:
    document = Document()
    document.add_heading("第一章 模板章节", level=1)
    document.add_paragraph(
        "这是学校模板中的正文示例段落，用于提取中文字体、字号、缩进、行距与对齐方式。"
    )
    document.add_paragraph("图1-1 模板结构图")
    document.add_paragraph("表1-1 模板数据")
    table = document.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "名称"
    table.cell(0, 1).text = "说明"
    table.cell(1, 0).text = "模板项"
    table.cell(1, 1).text = "模板值"
    document.save(path)


def build_source(path: Path) -> None:
    document = Document()
    document.add_heading("第一章 测试章节", level=1)
    document.add_paragraph(
        "这是论文原稿正文，CLI 只能修改它的格式，不能增删或改写任何可见文字。"
    )
    document.add_paragraph("图1-1 系统总体结构")
    document.add_paragraph("表1-1 测试数据")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "字段"
    table.cell(0, 1).text = "值"
    table.cell(1, 0).text = "库存"
    table.cell(1, 1).text = "128"
    document.save(path)


class FormatCliTests(unittest.TestCase):
    def test_published_hard_link_survives_staging_cleanup_failure(self) -> None:
        with tempfile.TemporaryDirectory(prefix="dokiai_publish_") as directory:
            root = Path(directory)
            staging = root / "working.docx"
            output = root / "output.docx"
            staging.write_bytes(b"complete-result")

            with patch.object(Path, "unlink", side_effect=PermissionError("busy")):
                _publish_without_overwrite(staging, output)

            self.assertEqual(output.read_bytes(), b"complete-result")

    def test_windows_rename_fallback_when_hard_links_are_unavailable(self) -> None:
        with tempfile.TemporaryDirectory(prefix="dokiai_publish_fallback_") as directory:
            root = Path(directory)
            staging = root / "working.docx"
            output = root / "output.docx"
            staging.write_bytes(b"complete-result")

            with patch("format_cli.os.link", side_effect=OSError("not supported")), patch(
                "format_cli.os.name", "nt"
            ):
                _publish_without_overwrite(staging, output)

            self.assertFalse(staging.exists())
            self.assertEqual(output.read_bytes(), b"complete-result")

    def test_body_is_editable_except_for_locked_first_line_indent(self) -> None:
        rules = DocumentRules()
        editable = _editable_rules(rules)
        body = editable["body"]["normal"]
        self.assertEqual(body["fontSizeName"], "小四")
        self.assertEqual(body["fontSizePt"], 12.0)
        self.assertEqual(body["fixedLineSpacingPt"], 20.0)
        self.assertEqual(body["multipleLineSpacing"], 1.25)

        with tempfile.TemporaryDirectory(prefix="dokiai_confirmed_rules_") as directory:
            confirmed = Path(directory) / "rules.json"
            confirmed.write_text(
                json.dumps(
                    {
                        "body": {
                            "normal": {
                                "chineseFont": "仿宋",
                                "latinFont": "Arial",
                                "fontSizePt": 14,
                                "lineSpacingMode": "multiple",
                                "multipleLineSpacing": 1.75,
                                "fixedLineSpacingPt": 24,
                                "spaceBefore": {"unit": "pt", "value": 6},
                                "spaceAfter": {"unit": "line", "value": 0.5},
                                "firstLineIndentChars": 0,
                            }
                        }
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            _apply_confirmed_rules(rules, confirmed)
            enforce_locked_document_policy(rules)

        body_rule = rules.normal_text
        self.assertEqual(body_rule.chinese_font, "仿宋")
        self.assertEqual(body_rule.latin_font, "Arial")
        self.assertEqual(body_rule.font_size_pt, 14.0)
        self.assertEqual(body_rule.line_spacing_mode, "multiple")
        self.assertEqual(body_rule.multiple_line_spacing, 1.75)
        self.assertEqual(body_rule.fixed_line_spacing_pt, 24.0)
        self.assertEqual(body_rule.space_before_pt, 6.0)
        self.assertEqual(body_rule.space_after_lines, 0.5)
        self.assertEqual(body_rule.special_indent_mode, "first_line")
        self.assertEqual(body_rule.special_indent_chars, 2.0)

    def run_cli(
        self,
        source: Path,
        template: Path,
        output: Path,
        result_json: Path,
        instructions: Path | None = None,
        use_doubao: bool = False,
    ) -> subprocess.CompletedProcess[str]:
        command = [
            sys.executable,
            "-B",
            str(CLI),
            "--source",
            str(source),
            "--template",
            str(template),
            "--output",
            str(output),
            "--result-json",
            str(result_json),
        ]
        if instructions is not None:
            command.extend(["--instructions-file", str(instructions)])
        if use_doubao:
            command.append("--use-doubao")
        environment = dict(os.environ)
        environment["PYTHONUTF8"] = "1"
        environment["PYTHONDONTWRITEBYTECODE"] = "1"
        return subprocess.run(
            command,
            cwd=TOOL_ROOT,
            env=environment,
            text=True,
            encoding="utf-8",
            capture_output=True,
            timeout=60,
            check=False,
        )

    def test_cli_formats_copy_and_preserves_all_visible_content(self) -> None:
        with tempfile.TemporaryDirectory(prefix="dokiai_format_cli_") as directory:
            root = Path(directory)
            source = root / "source.docx"
            template = root / "template.docx"
            output = root / "result.docx"
            result_json = root / "result.json"
            instructions = root / "instructions.txt"
            build_source(source)
            build_template(template)
            instructions.write_text(
                "正文宋体小四、1.5倍行距；表格仿宋二号、全框线、左对齐、"
                "首行缩进2字符。",
                encoding="utf-8",
            )
            source_before = file_hash(source)
            template_before = file_hash(template)
            text_before = visible_text(source)

            completed = self.run_cli(
                source, template, output, result_json, instructions
            )

            self.assertEqual(completed.returncode, 0, completed.stderr or completed.stdout)
            events = [json.loads(line) for line in completed.stdout.splitlines() if line]
            self.assertTrue(events)
            self.assertTrue(all(event["type"] == "progress" for event in events))
            self.assertIn("正在", events[0]["message"])
            self.assertEqual(events[-1]["progress"], 100)
            self.assertEqual(events[-1]["stage"], "completed")
            self.assertEqual([event["progress"] for event in events], sorted(event["progress"] for event in events))

            payload = json.loads(result_json.read_text(encoding="utf-8"))
            self.assertTrue(payload["success"])
            self.assertGreater(payload["changedCount"], 0)
            self.assertTrue(payload["integrity"]["passed"])
            self.assertEqual(payload["analysis"]["tableCount"], 1)
            table_rule = payload["ruleSummary"]["table"]
            self.assertEqual(table_rule["borderStyle"], "three_line")
            self.assertEqual(table_rule["chineseFont"], "宋体")
            self.assertEqual(table_rule["latinFont"], "Times New Roman")
            self.assertEqual(table_rule["fontSizeName"], "小四")
            self.assertEqual(table_rule["fontSizePt"], 12.0)
            self.assertFalse(table_rule["bold"])
            self.assertEqual(table_rule["alignment"], "center")
            self.assertEqual(table_rule["firstLineIndentChars"], 0.0)
            self.assertEqual(table_rule["verticalAlignment"], "center")
            self.assertFalse(table_rule["headerRowBold"])
            self.assertTrue(
                any("系统固定规范" in note for note in payload["instructionNotes"])
            )
            self.assertIsNone(payload["error"])
            self.assertTrue(output.is_file())
            output_text = visible_text(output)
            self.assertIn("目录", output_text)
            cursor = 0
            for item in text_before:
                cursor = output_text.index(item, cursor) + 1
            self.assertEqual(file_hash(source), source_before)
            self.assertEqual(file_hash(template), template_before)

    def test_doubao_result_cannot_override_locked_table_policy(self) -> None:
        with tempfile.TemporaryDirectory(prefix="dokiai_format_cli_") as directory:
            root = Path(directory)
            source = root / "source.docx"
            template = root / "template.docx"
            output = root / "result.docx"
            result_json = root / "result.json"
            instructions = root / "instructions.txt"
            build_source(source)
            build_template(template)
            instructions.write_text("请修改表格格式。", encoding="utf-8")

            def fake_doubao_parse(parser, requirement, rules):
                self.assertTrue(requirement)
                rules.table.enabled = False
                rules.table.border_style = "grid"
                rules.table.chinese_font = "仿宋"
                rules.table.latin_font = "Arial"
                rules.table.number_font = "Arial"
                rules.table.font_size_name = "二号"
                rules.table.font_size_pt = 22.0
                rules.table.bold = True
                rules.table.header_row_bold = True
                rules.table.alignment = "left"
                rules.table.special_indent_mode = "first_line"
                rules.table.special_indent_chars = 2.0
                rules.table.first_line_indent_chars = 2.0
                return rules, ["模拟豆包返回冲突表格规则"]

            args = Namespace(
                source=str(source),
                template=str(template),
                output=str(output),
                result_json=str(result_json),
                instructions_file=str(instructions),
                use_doubao=True,
            )
            with patch.object(
                format_cli.DoubaoRuleParser, "parse", new=fake_doubao_parse
            ):
                payload = run_job(args)

            table_rule = payload["ruleSummary"]["table"]
            self.assertTrue(table_rule["enabled"])
            self.assertEqual(table_rule["borderStyle"], "three_line")
            self.assertEqual(table_rule["chineseFont"], "宋体")
            self.assertEqual(table_rule["latinFont"], "Times New Roman")
            self.assertEqual(table_rule["fontSizeName"], "小四")
            self.assertEqual(table_rule["fontSizePt"], 12.0)
            self.assertFalse(table_rule["bold"])
            self.assertFalse(table_rule["headerRowBold"])
            self.assertEqual(table_rule["alignment"], "center")
            self.assertEqual(table_rule["firstLineIndentChars"], 0.0)
            self.assertTrue(payload["integrity"]["passed"])
            self.assertTrue(output.is_file())

    def test_existing_output_is_never_overwritten_and_failure_json_is_written(self) -> None:
        with tempfile.TemporaryDirectory(prefix="dokiai_format_cli_") as directory:
            root = Path(directory)
            source = root / "source.docx"
            template = root / "template.docx"
            output = root / "result.docx"
            result_json = root / "result.json"
            build_source(source)
            build_template(template)
            sentinel = b"do-not-overwrite"
            output.write_bytes(sentinel)

            completed = self.run_cli(source, template, output, result_json)

            self.assertNotEqual(completed.returncode, 0)
            self.assertEqual(output.read_bytes(), sentinel)
            payload = json.loads(result_json.read_text(encoding="utf-8"))
            self.assertFalse(payload["success"])
            self.assertEqual(payload["errorCode"], "OUTPUT_EXISTS")
            self.assertIn("拒绝覆盖", payload["error"])
            events = [json.loads(line) for line in completed.stdout.splitlines() if line]
            self.assertEqual(events[-1]["stage"], "failed")

    def test_result_json_alias_never_overwrites_an_input_document(self) -> None:
        with tempfile.TemporaryDirectory(prefix="dokiai_format_cli_") as directory:
            root = Path(directory)
            source = root / "source.docx"
            template = root / "template.docx"
            output = root / "result.docx"
            build_source(source)
            build_template(template)
            source_before = source.read_bytes()

            completed = self.run_cli(source, template, output, source)

            self.assertNotEqual(completed.returncode, 0)
            self.assertEqual(source.read_bytes(), source_before)
            self.assertFalse(output.exists())
            events = [json.loads(line) for line in completed.stdout.splitlines() if line]
            self.assertEqual(events[-1]["stage"], "failed")
            self.assertIn("路径必须互不相同", events[-1]["message"])

            instructions = root / "instructions.txt"
            instructions.write_text("正文宋体小四。", encoding="utf-8")
            instructions_before = instructions.read_bytes()
            second_output = root / "second-result.docx"
            second = self.run_cli(
                source, template, second_output, instructions, instructions
            )

            self.assertNotEqual(second.returncode, 0)
            self.assertEqual(instructions.read_bytes(), instructions_before)
            self.assertFalse(second_output.exists())

    def test_legacy_template_has_clear_non_windows_error(self) -> None:
        with tempfile.TemporaryDirectory(prefix="dokiai_format_cli_") as directory:
            template = Path(directory) / "legacy.doc"
            template.write_bytes(b"legacy")
            with self.assertRaisesRegex(
                WordConversionError,
                r"Linux/macOS.*另存为 \.docx",
            ):
                validate_runtime_support(template, platform_name="posix")

    def test_doubao_flag_without_extra_instructions_still_formats_template(self) -> None:
        with tempfile.TemporaryDirectory(prefix="dokiai_format_cli_") as directory:
            root = Path(directory)
            source = root / "source.docx"
            template = root / "template.docx"
            output = root / "result.docx"
            result_json = root / "result.json"
            build_source(source)
            build_template(template)

            completed = self.run_cli(
                source, template, output, result_json, use_doubao=True
            )

            self.assertEqual(completed.returncode, 0, completed.stderr or completed.stdout)
            self.assertTrue(output.exists())
            payload = json.loads(result_json.read_text(encoding="utf-8"))
            self.assertTrue(payload["success"])
            self.assertIsNone(payload["error"])

    def test_result_write_failure_removes_published_output(self) -> None:
        with tempfile.TemporaryDirectory(prefix="dokiai_format_cli_") as directory:
            root = Path(directory)
            source = root / "source.docx"
            template = root / "template.docx"
            output = root / "result.docx"
            result_json = root / "result.json"
            build_source(source)
            build_template(template)
            real_write_json = format_cli._write_json_atomic

            def fail_success_report(path: Path, payload: dict) -> None:
                if payload.get("success"):
                    raise OSError("simulated result report failure")
                real_write_json(path, payload)

            args = Namespace(
                source=str(source),
                template=str(template),
                output=str(output),
                result_json=str(result_json),
                instructions_file=None,
                use_doubao=False,
            )
            with patch.object(
                format_cli, "_write_json_atomic", side_effect=fail_success_report
            ):
                with self.assertRaisesRegex(OSError, "simulated"):
                    run_job(args)

            self.assertFalse(output.exists())
            self.assertFalse(output.with_suffix(".log.json").exists())
            payload = json.loads(result_json.read_text(encoding="utf-8"))
            self.assertFalse(payload["success"])


if __name__ == "__main__":
    unittest.main()
