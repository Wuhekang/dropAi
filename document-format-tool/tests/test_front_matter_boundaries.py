from __future__ import annotations

from pathlib import Path
import sys
import tempfile
import unittest

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))
from word_formatter.core.processor import DocumentProcessor
from word_formatter.core.template_text import _content_boundary, read_template_text
from word_formatter.models.results import ProcessResult


class FrontMatterBoundaryTests(unittest.TestCase):
    def test_later_introduction_cannot_move_retained_start_past_real_body(self) -> None:
        with tempfile.TemporaryDirectory(prefix="word_keep_chapters_") as directory:
            root = Path(directory)
            source_path, template_path = root / "source.docx", root / "template.docx"
            source = Document()
            source.add_heading("第一章 文献综述", 1)
            source.add_paragraph("作者第一章正文必须保留。")
            source.add_heading("第二章 研究设计", 1)
            source.add_paragraph("引言")
            source.add_paragraph("作者第二章正文必须保留。")
            source.save(source_path)
            template = Document()
            template.add_paragraph("学校新封面")
            template.add_paragraph("目录")
            template.save(template_path)
            fallback = DocumentProcessor._main_content_start(source)
            self.assertEqual(fallback, 1)
            self.assertEqual(DocumentProcessor._retained_content_start(source, fallback), 1)
            composed = DocumentProcessor._compose_with_template_front(
                source_path, template_path, ProcessResult(source_path, root / "out.docx"),
                {"documentKind": "template", "copyFrontMatter": True,
                 "frontMatterRange": {"startParagraph": 1, "endParagraph": 1}},
            )
            texts = [paragraph.text for paragraph in composed.paragraphs]
            self.assertEqual(texts[0], "学校新封面")
            for paragraph in source.paragraphs:
                self.assertIn(paragraph.text, texts)

    def test_unstructured_body_is_not_deleted_when_a_late_preface_is_found(self) -> None:
        source = Document()
        source.add_paragraph("这是一篇没有使用标题样式的原稿正文。")
        source.add_paragraph("正文后才出现的小节")
        source.add_paragraph("前言")
        self.assertEqual(DocumentProcessor._retained_content_start(source, 1), 1)

    def test_late_toc_does_not_hide_the_earliest_real_heading(self) -> None:
        source = Document()
        source.add_heading("第一章 文献综述", 1)
        source.add_paragraph("作者正文")
        source.styles.add_style("TOC 1", WD_STYLE_TYPE.PARAGRAPH)
        source.add_paragraph("缓存目录", style="TOC 1")
        source.add_heading("第二章 研究设计", 1)
        self.assertEqual(DocumentProcessor._main_content_start(source), 4)
        self.assertEqual(DocumentProcessor._retained_content_start(source, 4), 1)

    def test_decorative_cover_number_does_not_hide_author_abstract(self) -> None:
        source = Document()
        source.add_paragraph("（2027）届毕业论文")
        source.add_paragraph("姓名：张三")
        source.add_paragraph("摘要")
        source.add_paragraph("作者摘要必须保留。")
        source.add_heading("第一章 文献综述", 1)
        fallback = DocumentProcessor._main_content_start(source)
        self.assertEqual(fallback, 3)
        self.assertEqual(DocumentProcessor._retained_content_start(source, fallback), 3)

    def test_annotated_contents_title_stops_cover_at_paragraph_twenty(self) -> None:
        with tempfile.TemporaryDirectory(prefix="word_front_boundary_") as directory:
            path = Path(directory) / "template.docx"
            document = Document()
            for index in range(1, 64):
                text = {
                    3: "毕业设计成果", 7: "题 目：", 9: "学生姓名：",
                    11: "专 业：", 13: "指导教师：",
                    21: "目   录（三号黑体字）", 23: "前言…………1",
                    40: "前 言（三号黑体字）", 41: "示例前言内容",
                    43: "1 XXXXXXXXX（三号黑体字）", 46: "示例正文内容",
                }.get(index, "")
                paragraph = document.add_paragraph(text)
                if index == 20:
                    paragraph._p.get_or_add_pPr().append(OxmlElement("w:sectPr"))
            document.save(path)
            context = read_template_text(path)
            self.assertEqual(context["copyCandidate"]["startParagraph"], 1)
            self.assertEqual(context["copyCandidate"]["endParagraph"], 20)
            for title in ("前 言（三号黑体字）", "摘 要（小四号宋体）", "Contents (三号黑体字)"):
                self.assertTrue(_content_boundary(title), title)
            for prose in ("前言（研究背景）", "目录…………1", "封面（三号黑体字）", "要求目录（三号黑体字）"):
                self.assertFalse(_content_boundary(prose), prose)

    def test_copy_cover_preserves_real_preface_and_skips_toc_styles_and_fields(self) -> None:
        with tempfile.TemporaryDirectory(prefix="word_source_preface_") as directory:
            root = Path(directory)
            template_path = root / "template.docx"
            template = Document()
            template.add_paragraph("新学校封面")
            template.add_paragraph("目录（三号黑体字）")
            template.add_paragraph("模板示例正文不得复制")
            template.save(template_path)
            for title in ("摘要：", "Abstract:", "前 言", "引言", "绪论（三号黑体字）"):
                with self.subTest(title=title):
                    source_path = root / "source.docx"
                    source = Document()
                    source.add_paragraph("原学校封面")
                    source.styles.add_style("TOC 1", WD_STYLE_TYPE.PARAGRAPH)
                    source.add_paragraph(title, style="TOC 1")
                    begin = OxmlElement("w:fldChar")
                    begin.set(qn("w:fldCharType"), "begin")
                    source.add_paragraph().add_run()._r.append(begin)
                    source.add_paragraph(title)
                    end = OxmlElement("w:fldChar")
                    end.set(qn("w:fldCharType"), "end")
                    source.add_paragraph().add_run()._r.append(end)
                    source.add_paragraph(title)
                    source.add_paragraph("作者真实前言第一段必须保留。")
                    source.add_paragraph("作者真实前言第二段必须保留。")
                    source.add_heading("第一章 绪论", level=1)
                    source.add_paragraph("作者真实正文必须保留。")
                    source.save(source_path)
                    composed = DocumentProcessor._compose_with_template_front(
                        source_path, template_path,
                        ProcessResult(source_path, root / "output.docx"),
                        {"documentKind": "template", "copyFrontMatter": True,
                         "frontMatterRange": {"startParagraph": 1, "endParagraph": 1}},
                    )
                    texts = [paragraph.text for paragraph in composed.paragraphs]
                    self.assertEqual(texts.count(title), 1)
                    self.assertIn("作者真实前言第一段必须保留。", texts)
                    self.assertIn("作者真实前言第二段必须保留。", texts)
                    self.assertIn("作者真实正文必须保留。", texts)
                    self.assertNotIn("原学校封面", texts)
                    self.assertNotIn("模板示例正文不得复制", texts)


if __name__ == "__main__":
    unittest.main()
