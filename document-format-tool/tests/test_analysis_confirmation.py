from argparse import Namespace
import json
from pathlib import Path
import sys
import tempfile
import unittest
from unittest.mock import patch

from docx import Document
from docx.enum.section import WD_SECTION_START

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))
from format_cli import _apply_confirmed_rules, _restore_analyzed_rules, run_job
from word_formatter.core.integrity import sha256_file
from word_formatter.core.processor import DocumentProcessor
from word_formatter.models.results import ProcessResult
from word_formatter.models.rules import DocumentRules


class AnalysisConfirmationTests(unittest.TestCase):
    def test_confirmed_snapshot_is_used_without_reextracting_or_reanalyzing(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            source, template = root / 'source.docx', root / 'template.docx'
            doc = Document()
            doc.add_paragraph('作者原封面')
            doc.add_heading('第一章 绪论', 1)
            doc.add_paragraph('这是论文正文，保留文字并按客户确认的字号排版。')
            doc.save(source)
            doc = Document()
            doc.add_paragraph('毕业论文撰写规范，封面见其他附件。')
            doc.save(template)
            rules = DocumentRules()
            rules.page_setup.margin_top_mm = 30
            rules.normal_text.chinese_font = '仿宋'
            rules.normal_text.line_spacing_mode = 'multiple'
            rules.normal_text.multiple_line_spacing = 1.25
            rules.heading_1.enabled = True
            rules.heading_1.font_size_pt = 18
            rules.heading_1.font_size_name = '小二'
            analysis = {'documentKind': 'specification', 'copyFrontMatter': False,
                        'frontMatterRange': None, 'reason': '封面另见附件', 'warnings': []}
            confirmed = root / 'confirmed.json'
            confirmed.write_text(json.dumps({
                'analyzedRules': rules.to_dict(), 'templateSha256': sha256_file(template),
                'templateAnalysis': analysis,
                'editableRules': {'body': {'normal': {'fontSizePt': 14, 'alignment': 'left', 'bold': False}}},
            }), encoding='utf-8')
            args = Namespace(source=str(source), template=str(template), output=str(root / 'out.docx'),
                             result_json=str(root / 'result.json'), instructions_file=None,
                             use_doubao=False, analyze_only=False, rules_file=str(confirmed))
            with patch('format_cli.TemplateRuleExtractor.extract', side_effect=AssertionError('must reuse snapshot')), \
                 patch('format_cli.DoubaoRuleParser.analyze_template', side_effect=AssertionError('must not rerun AI')), \
                 patch('word_formatter.core.word_converter.WordDocumentConverter.update_fields_in_place'):
                result = run_job(args)
            self.assertTrue(result['success'])
            self.assertEqual(result['ruleSummary']['pageSetup']['margin_top_mm'], 30)
            self.assertEqual(result['ruleSummary']['normalText']['chineseFont'], '仿宋')
            self.assertEqual(result['ruleSummary']['normalText']['fontSizeName'], '四号')
            self.assertEqual(result['ruleSummary']['heading1']['fontSizePt'], 18)
            output = Document(root / 'out.docx')
            text = '\n'.join(p.text for p in output.paragraphs)
            self.assertIn('作者原封面', text)
            self.assertNotIn('封面见其他附件', text)
            body = next(p for p in output.paragraphs if p.text.startswith('这是论文正文'))
            self.assertEqual(body.runs[0].font.size.pt, 14)

    def test_snapshot_rejects_changed_template(self):
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / 'rules.json'
            path.write_text(json.dumps({'analyzedRules': DocumentRules().to_dict(),
                                       'templateSha256': 'old', 'templateAnalysis': {}}), encoding='utf-8')
            with self.assertRaisesRegex(ValueError, '模板已变化'):
                _restore_analyzed_rules(path, 'changed')

    def test_copy_range_does_not_replace_authors_abstract_with_template_sample(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            template, source = root / 'template.docx', root / 'source.docx'
            doc = Document()
            doc.add_paragraph('学校封面')
            doc.add_paragraph('姓名：XXX 学号：XXX')
            doc.add_heading('摘要', 1)
            doc.add_paragraph('不能复制的模板摘要样例')
            doc.add_heading('第一章 模板正文', 1)
            doc.save(template)
            doc = Document()
            doc.add_paragraph('原稿旧封面')
            doc.add_heading('摘要：', 1)
            doc.add_paragraph('作者自己的摘要必须保留')
            doc.styles.add_style('TOC 1', 1)
            doc.add_paragraph('第一章 作者正文\t1', style='TOC 1')
            doc.add_heading('第一章 作者正文', 1)
            doc.save(source)
            policy = {'documentKind': 'mixed', 'copyFrontMatter': True,
                      'frontMatterRange': {'startParagraph': 1, 'endParagraph': 2}}
            composed = DocumentProcessor._compose_with_template_front(
                source, template, ProcessResult(source, root / 'out.docx'), policy)
            text = '\n'.join(p.text for p in composed.paragraphs)
            self.assertIn('学校封面', text)
            self.assertIn('作者自己的摘要必须保留', text)
            self.assertNotIn('模板摘要样例', text)
            self.assertNotIn('原稿旧封面', text)
            self.assertTrue(composed.paragraphs[1]._p.xpath('./w:pPr/w:sectPr'))

    def test_toc_does_not_double_existing_section_boundary_on_refresh(self):
        document = Document()
        document.add_paragraph('封面')
        document.add_section(WD_SECTION_START.NEW_PAGE)
        document.add_heading('第一章 正文', 1)
        result = ProcessResult(Path('source'), Path('out'))
        rules = DocumentRules()
        DocumentProcessor._ensure_toc(document, rules, result, 3)
        DocumentProcessor._apply_toc(document, rules, result)
        title = next(p for p in document.paragraphs if p.text == '目录')
        self.assertFalse(title.paragraph_format.page_break_before)

    def test_copied_cover_table_stays_with_title_and_section_follows_table(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            template, source = root / 'template.docx', root / 'source.docx'
            doc = Document()
            doc.add_paragraph('学校封面标题')
            doc.add_table(rows=1, cols=1).cell(0, 0).text = '封面姓名与学号'
            doc.add_heading('摘要', 1)
            doc.save(template)
            doc = Document()
            doc.add_heading('摘要', 1)
            doc.add_paragraph('作者摘要')
            doc.save(source)
            composed = DocumentProcessor._compose_with_template_front(
                source, template, ProcessResult(source, root / 'out.docx'),
                {'documentKind': 'template', 'copyFrontMatter': True,
                 'frontMatterRange': {'startParagraph': 1, 'endParagraph': 1}})
            self.assertFalse(composed.paragraphs[0]._p.xpath('./w:pPr/w:sectPr'))
            after_table = composed.tables[0]._tbl.getnext()
            self.assertTrue(after_table.xpath('./w:pPr/w:sectPr'))

    def test_numeric_confirmation_is_finite_and_supports_points_over_twenty(self):
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / 'rules.json'
            rule = {'spaceBefore': {'unit': 'pt', 'value': 24}}
            path.write_text(json.dumps({'body': {'normal': rule}}), encoding='utf-8')
            rules = DocumentRules()
            _apply_confirmed_rules(rules, path)
            self.assertEqual(rules.normal_text.space_before_pt, 24)
            rule['fontSizePt'] = float('nan')
            path.write_text(json.dumps({'body': {'normal': rule}}), encoding='utf-8')
            with self.assertRaises(ValueError):
                _apply_confirmed_rules(rules, path)


if __name__ == '__main__':
    unittest.main()
