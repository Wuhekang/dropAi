from argparse import Namespace
import json
from pathlib import Path
import sys
import tempfile
import unittest
from unittest.mock import patch

from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))
from format_cli import run_job, _format_report
from word_formatter.core.integrity import IntegrityValidationError, sha256_file
from word_formatter.models.results import ChangeRecord, ProcessResult
from word_formatter.models.rules import DocumentRules


class CliDeliveryPolicyTests(unittest.TestCase):
    def setUp(self):
        self.temp = tempfile.TemporaryDirectory()
        self.addCleanup(self.temp.cleanup)
        self.root = Path(self.temp.name)
        self.source, self.template = self.root / 'source.docx', self.root / 'template.docx'
        for path in (self.source, self.template):
            doc = Document()
            doc.add_heading('第一章 绪论', 1)
            doc.add_paragraph('作者正文。见第1页。')
            doc.save(path)
        self.original_hash = sha256_file(self.source)
        confirmed = self.root / 'confirmed.json'
        confirmed.write_text(json.dumps({
            'analyzedRules': DocumentRules().to_dict(),
            'templateSha256': sha256_file(self.template),
            'templateAnalysis': {'documentKind': 'specification', 'copyFrontMatter': False},
            'editableRules': {},
        }), encoding='utf-8')
        self.args = Namespace(source=str(self.source), template=str(self.template),
            output=str(self.root / 'output.docx'), result_json=str(self.root / 'result.json'),
            rules_file=str(confirmed), instructions_file=None, use_doubao=False, analyze_only=False)

    def fake_processor(self, source, rules, output, template, **kwargs):
        # Simulate a permitted refreshed field result; the detailed check must
        # honestly report a mismatch, while the finished file remains available.
        doc = Document(source)
        doc.paragraphs[1].text = '作者正文。见第2页。'
        doc.save(output)
        result = ProcessResult(source, output)
        result.records.append(ChangeRecord(2, '普通正文', '原格式', '确认格式', '已执行格式规则'))
        return result

    def test_text_difference_returns_downloadable_file_and_honest_warning(self):
        with patch('format_cli.DocumentProcessor.process', side_effect=self.fake_processor):
            payload = run_job(self.args)
        self.assertTrue(payload['success'])
        self.assertFalse(payload['integrity']['passed'])
        self.assertEqual(payload['integrity']['mode'], 'format_first')
        self.assertTrue(payload['integrity']['basicChecksPassed'])
        self.assertTrue(payload['integrity']['deliveryAllowed'])
        self.assertIn('body_text_sha256', payload['integrity']['differences'])
        self.assertTrue(any('文字对比' in value for value in payload['warnings']))
        self.assertIn({'item': '普通正文', 'count': 1}, payload['formatReport']['applied'])
        self.assertTrue(payload['formatReport']['notApplied'])
        self.assertTrue(Path(self.args.output).is_file())
        self.assertEqual(Document(self.args.output).paragraphs[1].text, '作者正文。见第2页。')
        self.assertEqual(sha256_file(self.source), self.original_hash)

    def test_unreadable_output_is_not_published(self):
        def broken(source, rules, output, template, **kwargs):
            output.write_bytes(b'not a Word document')
            return ProcessResult(source, output)
        with patch('format_cli.DocumentProcessor.process', side_effect=broken):
            with self.assertRaises(IntegrityValidationError):
                run_job(self.args)
        self.assertFalse(Path(self.args.output).exists())
        self.assertFalse(json.loads(Path(self.args.result_json).read_text(encoding='utf-8'))['success'])

    def test_changed_source_is_still_rejected(self):
        def changing(source, rules, output, template, **kwargs):
            result = self.fake_processor(source, rules, output, template, **kwargs)
            doc = Document(source)
            doc.add_paragraph('意外改写原稿')
            doc.save(source)
            return result
        with patch('format_cli.DocumentProcessor.process', side_effect=changing):
            with self.assertRaisesRegex(IntegrityValidationError, '源文件发生变化'):
                run_job(self.args)
        self.assertFalse(Path(self.args.output).exists())

    def test_report_deduplicates_toc_passes_and_does_not_claim_skips_succeeded(self):
        result = ProcessResult(self.source, self.root / 'out.docx')
        result.records = [
            ChangeRecord(2, '目录标题', '旧', '新', '已处理'),
            ChangeRecord(2, '目录标题', '旧', '新', '二次刷新'),
            ChangeRecord(3, '复杂对象', '旧', '旧', '不支持自动处理', status='skipped'),
        ]
        report = _format_report(result, DocumentRules(), [])
        self.assertEqual(report['applied'], [{'item': '目录标题', 'count': 1}])
        self.assertIn({'item': '复杂对象', 'reason': '不支持自动处理'}, report['notApplied'])


if __name__ == '__main__':
    unittest.main()
