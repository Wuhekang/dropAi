from argparse import Namespace
from dataclasses import fields
import json
from pathlib import Path
import shutil
import sys
import tempfile
from types import SimpleNamespace
import unittest
from unittest.mock import patch

from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))
from format_cli import _apply_confirmed_rules, _forward_template_progress, run_job
from word_formatter.core.integrity import sha256_file
from word_formatter.models.results import ChangeRecord, ProcessResult
from word_formatter.models.rules import (
    DEFAULT_LATIN_FONT, DocumentRules, ParagraphRule, apply_default_latin_fonts,
    enforce_locked_document_policy,
)


class LatinFontDefaultTests(unittest.TestCase):
    def setUp(self):
        self.temp = tempfile.TemporaryDirectory()
        self.addCleanup(self.temp.cleanup)
        self.root = Path(self.temp.name)
        self.source = self.root / "source.docx"
        self.template = self.root / "template.docx"
        for path in (self.source, self.template):
            document = Document()
            document.add_heading("第一章 正文", 1)
            document.add_paragraph("中文 English 123")
            document.save(path)
        self.analysis = {
            "documentKind": "specification", "copyFrontMatter": False,
            "frontMatterRange": None, "reason": "文字撰写规范", "warnings": [],
        }

    @staticmethod
    def inferred_rules():
        rules = DocumentRules()
        for field in fields(rules):
            rule = getattr(rules, field.name)
            if isinstance(rule, ParagraphRule):
                rule.latin_font = "黑体"
                rule.number_font = "宋体"
        rules.normal_text.chinese_font = "仿宋"
        return rules

    def args(self, name, *, analyze_only=False, rules_file=None, instructions=None, use_doubao=False):
        instruction_path = None
        if instructions:
            instruction_path = self.root / (name + "-instructions.txt")
            instruction_path.write_text(instructions, encoding="utf-8")
        return Namespace(source=str(self.source), template=str(self.template),
                         output=str(self.root / (name + ".docx")),
                         result_json=str(self.root / (name + ".json")),
                         instructions_file=str(instruction_path) if instruction_path else None,
                         rules_file=str(rules_file) if rules_file else None,
                         analyze_only=analyze_only, use_doubao=use_doubao)

    @staticmethod
    def process_without_word(source, rules, output, template, **kwargs):
        shutil.copyfile(source, output)
        result = ProcessResult(source, output)
        result.records.append(ChangeRecord(2, "普通正文", "旧格式", "新格式", "测试格式规则交接"))
        return result

    def test_all_inferred_paragraph_fonts_default_to_times_new_roman(self):
        rules = self.inferred_rules()
        apply_default_latin_fonts(rules)
        enforce_locked_document_policy(rules)
        for field in fields(rules):
            rule = getattr(rules, field.name)
            if isinstance(rule, ParagraphRule):
                self.assertEqual(rule.latin_font, DEFAULT_LATIN_FONT, field.name)
                self.assertEqual(rule.number_font, DEFAULT_LATIN_FONT, field.name)
        self.assertEqual(rules.normal_text.chinese_font, "仿宋")
        self.assertEqual(rules.table.chinese_font, "宋体")

    def test_local_template_and_natural_language_parser_cannot_change_english_default(self):
        def local_instructions(text, rules):
            rules.normal_text.latin_font = "Arial"
            rules.heading_1.latin_font = "黑体"
            return []

        with patch("format_cli.TemplateRuleExtractor.extract", return_value=SimpleNamespace(rules=self.inferred_rules(), notes=[])), \
             patch("format_cli.NaturalLanguageRuleParser.apply", side_effect=local_instructions), \
             patch("format_cli.DocumentProcessor.process", side_effect=self.process_without_word):
            payload = run_job(self.args("local", instructions="一级标题居中"))
        self.assertEqual(payload["ruleSummary"]["normalText"]["latinFont"], DEFAULT_LATIN_FONT)
        self.assertEqual(payload["ruleSummary"]["heading1"]["latinFont"], DEFAULT_LATIN_FONT)
        self.assertEqual(payload["ruleSummary"]["table"]["latinFont"], DEFAULT_LATIN_FONT)

    def test_ai_success_and_timeout_fallback_have_the_same_defaults(self):
        for fallback in (False, True):
            with self.subTest(fallback=fallback):
                rules = self.inferred_rules()
                analysis = {**self.analysis, "warnings": ["AI 超时，保留程序识别值"] if fallback else []}
                parser = SimpleNamespace(last_template_analysis=analysis,
                                         analyze_template=lambda *args: (rules, []))
                with patch("format_cli.TemplateRuleExtractor.extract", return_value=SimpleNamespace(rules=rules, notes=[])), \
                     patch("format_cli.DoubaoRuleParser", return_value=parser):
                    payload = run_job(self.args("fallback" if fallback else "ai", analyze_only=True))
                for name, rule in payload["analyzedRules"].items():
                    if isinstance(rule, dict) and "latin_font" in rule:
                        self.assertEqual(rule["latin_font"], DEFAULT_LATIN_FONT, name)
                        self.assertEqual(rule["number_font"], DEFAULT_LATIN_FONT, name)
                if fallback:
                    self.assertIn("AI 超时，保留程序识别值", payload["warnings"])

    def test_additional_ai_parser_cannot_override_english_default(self):
        rules = self.inferred_rules()
        parser = SimpleNamespace(parse=lambda *args: (rules, []))
        with patch("format_cli.TemplateRuleExtractor.extract", return_value=SimpleNamespace(rules=rules, notes=[])), \
             patch("format_cli.DoubaoRuleParser", return_value=parser), \
             patch("format_cli.DocumentProcessor.process", side_effect=self.process_without_word):
            payload = run_job(self.args("additional-ai", instructions="图名居中", use_doubao=True))
        self.assertEqual(payload["ruleSummary"]["figureCaption"]["latinFont"], DEFAULT_LATIN_FONT)

    def test_analysis_snapshot_keeps_explicit_customer_english_font(self):
        rules = self.inferred_rules()
        parser = SimpleNamespace(last_template_analysis=self.analysis,
                                 analyze_template=lambda *args: (rules, []))
        with patch("format_cli.TemplateRuleExtractor.extract", return_value=SimpleNamespace(rules=rules, notes=[])), \
             patch("format_cli.DoubaoRuleParser", return_value=parser):
            analyzed = run_job(self.args("analyzed", analyze_only=True))
        confirmed = self.root / "confirmed.json"
        confirmed.write_text(json.dumps({
            "analyzedRules": analyzed["analyzedRules"], "templateSha256": analyzed["templateSha256"],
            "templateAnalysis": analyzed["templateAnalysis"],
            "editableRules": {"body": {"normal": {"latinFont": "Arial"}},
                              "headings": {"level1": {"latinFont": "Georgia"}}},
        }), encoding="utf-8")

        def inspect_rules(source, rules, output, template, **kwargs):
            self.assertEqual(rules.normal_text.latin_font, "Arial")
            self.assertEqual(rules.normal_text.number_font, "Arial")
            self.assertEqual(rules.heading_1.latin_font, "Georgia")
            self.assertEqual(rules.heading_1.number_font, "Georgia")
            self.assertEqual(rules.heading_2.latin_font, DEFAULT_LATIN_FONT)
            return self.process_without_word(source, rules, output, template, **kwargs)

        with patch("format_cli.TemplateRuleExtractor.extract", side_effect=AssertionError("must reuse snapshot")), \
             patch("format_cli.DoubaoRuleParser", side_effect=AssertionError("must not call AI")), \
             patch("format_cli.DocumentProcessor.process", side_effect=inspect_rules):
            result = run_job(self.args("confirmed", rules_file=confirmed))
        self.assertEqual(result["ruleSummary"]["normalText"]["latinFont"], "Arial")
        self.assertEqual(result["ruleSummary"]["heading1"]["latinFont"], "Georgia")

    def test_legacy_snapshot_inherited_chinese_font_is_normalized_without_edit(self):
        confirmed = self.root / "old-confirmed.json"
        confirmed.write_text(json.dumps({
            "analyzedRules": self.inferred_rules().to_dict(), "templateSha256": sha256_file(self.template),
            "templateAnalysis": self.analysis, "editableRules": {"body": {"normal": {"fontSizePt": 14}}},
        }), encoding="utf-8")
        with patch("format_cli.DocumentProcessor.process", side_effect=self.process_without_word):
            result = run_job(self.args("legacy", rules_file=confirmed))
        self.assertEqual(result["ruleSummary"]["normalText"]["latinFont"], DEFAULT_LATIN_FONT)
        self.assertEqual(result["ruleSummary"]["normalText"]["fontSizePt"], 14)

    def test_unchanged_legacy_form_font_is_default_but_changed_font_is_explicit(self):
        rules = self.inferred_rules()
        snapshot = rules.to_dict()
        apply_default_latin_fonts(rules)
        confirmed = self.root / "legacy-form.json"
        confirmed.write_text(json.dumps({
            "analyzedRules": snapshot,
            "editableRules": {"body": {"normal": {"latinFont": "黑体"}},
                              "headings": {"level1": {"latinFont": "Arial"}},
                              "captions": {"figure": {"latinFont": "黑体"}}},
        }), encoding="utf-8")
        _apply_confirmed_rules(rules, confirmed)
        self.assertEqual(rules.normal_text.latin_font, DEFAULT_LATIN_FONT)
        self.assertEqual(rules.normal_text.number_font, DEFAULT_LATIN_FONT)
        self.assertEqual(rules.figure_caption.latin_font, DEFAULT_LATIN_FONT)
        self.assertEqual(rules.heading_1.latin_font, "Arial")
        self.assertEqual(rules.heading_1.number_font, "Arial")

    def test_confirmation_without_snapshot_keeps_explicit_font(self):
        rules = DocumentRules()
        confirmed = self.root / "no-snapshot.json"
        confirmed.write_text(json.dumps({"body": {"normal": {"latinFont": "Arial"}}}), encoding="utf-8")
        _apply_confirmed_rules(rules, confirmed)
        self.assertEqual(rules.normal_text.latin_font, "Arial")
        self.assertEqual(rules.normal_text.number_font, "Arial")

    def test_request_completion_progress_does_not_claim_recognition_succeeded(self):
        with patch("format_cli.emit_progress") as emit:
            _forward_template_progress({"stage": "map", "completed": 0, "total": 8, "successful": 0})
            _forward_template_progress({"stage": "map", "completed": 4, "total": 8, "successful": 0})
            _forward_template_progress({"stage": "map", "completed": 8, "total": 8, "successful": 0})
            _forward_template_progress({"stage": "reduce", "completed": 8, "total": 8, "successful": 0})
        self.assertEqual([call.args[0] for call in emit.call_args_list], [24, 26, 28, 29])
        for call in emit.call_args_list:
            self.assertEqual(call.args[1], "analyzing_template")
            self.assertNotIn("识别成功", call.args[2])
            self.assertNotIn("分析成功", call.args[2])
        self.assertIn("4/8", emit.call_args_list[1].args[2])
        self.assertIn("整合", emit.call_args_list[-1].args[2])

    def test_cli_connects_pipeline_progress_callback_before_analysis(self):
        rules = self.inferred_rules()
        parser = SimpleNamespace(last_template_analysis=self.analysis)

        def analyze(*args):
            parser.template_progress_callback({"stage": "map", "completed": 3, "total": 4, "successful": 0})
            parser.template_progress_callback({"stage": "reduce", "completed": 4, "total": 4, "successful": 0})
            return rules, []

        parser.analyze_template = analyze
        with patch("format_cli.TemplateRuleExtractor.extract", return_value=SimpleNamespace(rules=rules, notes=[])), \
             patch("format_cli.DoubaoRuleParser", return_value=parser), \
             patch("format_cli.emit_progress") as emit:
            run_job(self.args("progress", analyze_only=True))
        notifications = [call.args for call in emit.call_args_list if call.args[1] == "analyzing_template"]
        self.assertEqual([item[0] for item in notifications], [24, 27, 29])
        self.assertIn("3/4", notifications[1][2])


if __name__ == "__main__":
    unittest.main()
