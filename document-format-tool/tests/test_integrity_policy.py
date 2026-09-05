from __future__ import annotations

from pathlib import Path
import shutil
import sys
import tempfile
import unittest
from zipfile import ZipFile

from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))
from word_formatter.core.integrity import (
    IntegrityValidationError,
    inspect_docx,
    sha256_file,
    validate_preservation,
)


class IntegrityPolicyTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temporary = tempfile.TemporaryDirectory(prefix="word_integrity_policy_")
        self.addCleanup(self.temporary.cleanup)
        self.root = Path(self.temporary.name)
        self.source = self.root / "source.docx"
        self.output = self.root / "output.docx"
        document = Document()
        document.add_paragraph("原文第一段。")
        document.add_paragraph("原文第二段。")
        document.add_table(rows=1, cols=1).cell(0, 0).text = "原表格内容"
        document.save(self.source)
        self.original_hash = sha256_file(self.source)
        shutil.copyfile(self.source, self.output)

    def replace_part(self, name: str, content: bytes | None) -> None:
        with ZipFile(self.output) as package:
            parts = {info.filename: package.read(info.filename) for info in package.infolist()}
        if content is None:
            parts.pop(name)
        else:
            parts[name] = content
        with ZipFile(self.output, "w") as package:
            for part_name, data in parts.items():
                package.writestr(part_name, data)

    def changed_output(self) -> None:
        document = Document()
        document.add_paragraph("格式处理后有可记录的文字和结构差异。")
        document.save(self.output)

    def test_format_first_delivers_but_retains_text_and_count_differences(self) -> None:
        self.changed_output()
        result = validate_preservation(
            self.source, self.output,
            expected_source_sha256=self.original_hash,
            strict=False,
        )
        summary = result.summary()
        self.assertFalse(result.passed)
        self.assertFalse(summary["passed"])
        self.assertEqual(summary["mode"], "format_first")
        self.assertTrue(summary["basicChecksPassed"])
        self.assertTrue(summary["deliveryAllowed"])
        self.assertIn("body_text_sha256", result.differences)
        self.assertIn("paragraph_count", result.differences)
        self.assertEqual(result.differences["table_count"], {"before": 1, "after": 0})
        self.assertEqual(summary["differences"], result.differences)
        self.assertEqual(sha256_file(self.source), self.original_hash)

    def test_format_first_with_front_policy_retains_remaining_diagnostics(self) -> None:
        self.changed_output()
        result = validate_preservation(
            self.source, self.output, allow_front_matter=True, strict=False,
        )
        self.assertFalse(result.passed)
        self.assertIn("body_text_sha256", result.differences)
        self.assertIn("table_count", result.differences)
        self.assertTrue(result.summary()["deliveryAllowed"])

    def test_strict_remains_default_and_rejects_differences(self) -> None:
        self.changed_output()
        with self.assertRaisesRegex(IntegrityValidationError, "严格完整性校验未通过"):
            validate_preservation(self.source, self.output)
        with self.assertRaises(IntegrityValidationError):
            validate_preservation(self.source, self.output, strict=True)

    def test_unchanged_package_passes_both_modes(self) -> None:
        for strict in (False, True):
            with self.subTest(strict=strict):
                summary = validate_preservation(self.source, self.output, strict=strict).summary()
                self.assertTrue(summary["passed"])
                self.assertTrue(summary["basicChecksPassed"])
                self.assertTrue(summary["deliveryAllowed"])
                self.assertEqual(summary["differences"], {})
                self.assertEqual(summary["mode"], "strict" if strict else "format_first")

    def test_format_first_still_rejects_invalid_zip(self) -> None:
        self.output.write_bytes(b"not a DOCX ZIP archive")
        with self.assertRaisesRegex(IntegrityValidationError, "ZIP"):
            validate_preservation(self.source, self.output, strict=False)

    def test_format_first_still_rejects_missing_required_part(self) -> None:
        self.replace_part("[Content_Types].xml", None)
        with self.assertRaisesRegex(IntegrityValidationError, "缺少必要"):
            validate_preservation(self.source, self.output, strict=False)

    def test_format_first_still_rejects_malformed_xml_in_any_package_part(self) -> None:
        for part in ("word/document.xml", "word/styles.xml", "[Content_Types].xml"):
            with self.subTest(part=part):
                shutil.copyfile(self.source, self.output)
                self.replace_part(part, b"<malformed>")
                with self.assertRaisesRegex(IntegrityValidationError, "损坏的 OOXML"):
                    validate_preservation(self.source, self.output, strict=False)

    def test_default_input_inspection_does_not_add_extension_xml_restrictions(self) -> None:
        self.replace_part("word/repairableExtension.xml", b"<malformed>")
        self.assertIn("body_text_sha256", inspect_docx(self.output))
        with self.assertRaisesRegex(IntegrityValidationError, "损坏的 OOXML"):
            inspect_docx(self.output, validate_all_xml=True)
        with self.assertRaisesRegex(IntegrityValidationError, "损坏的 OOXML"):
            validate_preservation(self.source, self.output, strict=False)

    def test_format_first_still_rejects_missing_document_body(self) -> None:
        self.replace_part(
            "word/document.xml",
            b'<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>',
        )
        with self.assertRaisesRegex(IntegrityValidationError, "缺少正文"):
            validate_preservation(self.source, self.output, strict=False)

    def test_format_first_still_rejects_changed_source(self) -> None:
        source = Document(self.source)
        source.add_paragraph("处理期间外部修改了原稿。")
        source.save(self.source)
        with self.assertRaisesRegex(IntegrityValidationError, "源文件发生变化"):
            validate_preservation(
                self.source, self.output,
                expected_source_sha256=self.original_hash,
                strict=False,
            )


if __name__ == "__main__":
    unittest.main()
