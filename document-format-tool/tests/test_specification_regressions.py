from __future__ import annotations

from pathlib import Path
import shutil
import sys
import tempfile
import unittest
import zipfile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor


TOOL_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(TOOL_ROOT))

from word_formatter.core.integrity import validate_preservation  # noqa: E402
from word_formatter.core.finalizer import finalize_docx  # noqa: E402
from word_formatter.core.processor import DocumentProcessor  # noqa: E402
from word_formatter.core.template_extractor import TemplateRuleExtractor  # noqa: E402
from word_formatter.models.results import ProcessResult  # noqa: E402
from word_formatter.models.rules import DocumentRules  # noqa: E402


def build_written_specification(path: Path) -> None:
    document = Document()
    document.add_paragraph("附件3")
    document.add_paragraph("毕业设计（论文）撰写规范")
    table = document.add_table(rows=5, cols=2)
    values = [
        ("页面设置", "A4纸；正文中文宋体小四号，英文及数字Times New Roman小四号；两端对齐；首行缩进2字符；1.25倍行距。"),
        ("目录", "目录标题黑体小二号、居中、1.5倍行距、段前段后20磅；目录内容黑体五号、两端对齐、1.25倍行距。"),
        ("正文", "一级标题黑体小二号、居中、1.5倍行距、段前段后20磅；二级标题黑体小三号、居左顶格、1.5倍行距；三级标题黑体四号、居左顶格、1.5倍行距；四级标题黑体小四号、居左顶格、1.5倍行距；五级标题及正文宋体小四号、首行缩进2字符、1.25倍行距。"),
        ("图", "图名称宋体五号，英文及数字Times New Roman五号，1.25倍行距，图下方居中。"),
        ("表", "表名称宋体五号，英文及数字Times New Roman五号，1.25倍行距；要求表名称应标在表的上方、居中。"),
    ]
    for row, (label, value) in zip(table.rows, values):
        row.cells[0].text = label
        row.cells[1].text = value
    document.save(path)


class WrittenSpecificationTests(unittest.TestCase):
    def test_extracts_explicit_rules_without_shorter_font_size_collisions(self) -> None:
        with tempfile.TemporaryDirectory(prefix="dokiai_spec_") as directory:
            path = Path(directory) / "spec.docx"
            build_written_specification(path)

            extraction = TemplateRuleExtractor().extract(path)
            rules = extraction.rules

            self.assertEqual((rules.normal_text.font_size_name, rules.normal_text.font_size_pt), ("小四", 12.0))
            self.assertEqual((rules.heading_1.font_size_name, rules.heading_1.font_size_pt), ("小二", 18.0))
            self.assertEqual((rules.heading_2.font_size_name, rules.heading_2.font_size_pt), ("小三", 15.0))
            self.assertEqual(rules.heading_2.alignment, "left")
            self.assertEqual(rules.heading_1.space_before_pt, 20.0)
            self.assertEqual(rules.heading_1.space_after_pt, 20.0)
            self.assertEqual((rules.toc_title.font_size_name, rules.toc_title.font_size_pt), ("小二", 18.0))
            self.assertEqual((rules.toc_1.font_size_name, rules.toc_1.font_size_pt), ("五号", 10.5))
            self.assertEqual(rules.figure_caption.alignment, "center")
            self.assertEqual(rules.table_caption.alignment, "center")
            self.assertTrue(any("撰写规范表" in note for note in extraction.notes))

    def test_specification_text_is_not_copied_as_front_matter(self) -> None:
        with tempfile.TemporaryDirectory(prefix="dokiai_spec_front_") as directory:
            root = Path(directory)
            spec = root / "spec.docx"
            source = root / "source.docx"
            build_written_specification(spec)
            source_doc = Document()
            source_doc.add_paragraph("论文原封面")
            source_doc.add_heading("摘要", level=1)
            source_doc.add_paragraph("论文摘要正文")
            source_doc.save(source)
            result = ProcessResult(source, root / "output.docx")

            composed = DocumentProcessor._compose_with_template_front(source, spec, result)

            text = "\n".join(paragraph.text for paragraph in composed.paragraphs)
            self.assertIn("论文原封面", text)
            self.assertNotIn("毕业设计（论文）撰写规范", text)
            self.assertTrue(any("不复制" in warning for warning in result.warnings))

    def test_toc_starts_on_new_page_and_excludes_non_content_headings(self) -> None:
        document = Document()
        cover = document.add_heading("封面", level=1)
        document.add_heading("摘要", level=1)
        document.add_heading("第一章 绪论", level=1)
        appendix = document.add_heading("附录", level=1)
        references = document.add_heading("参考文献", level=1)
        acknowledgements = document.add_heading("致谢", level=1)
        result = ProcessResult(Path("source.docx"), Path("output.docx"))

        DocumentProcessor._exclude_non_content_toc_entries(document, 1)
        DocumentProcessor._ensure_toc(document, DocumentRules(), result, 2)

        toc_title = next(p for p in document.paragraphs if p.text == "目录")
        self.assertTrue(toc_title.paragraph_format.page_break_before)
        self.assertEqual(cover._p.pPr.find(qn("w:outlineLvl")).get(qn("w:val")), "9")
        self.assertEqual(appendix._p.pPr.find(qn("w:outlineLvl")).get(qn("w:val")), "9")
        self.assertIsNone(references._p.pPr.find(qn("w:outlineLvl")))
        self.assertIsNone(acknowledgements._p.pPr.find(qn("w:outlineLvl")))


class IntegrityFrontMatterTests(unittest.TestCase):
    def test_unreferenced_front_media_does_not_block_front_matter_replacement(self) -> None:
        with tempfile.TemporaryDirectory(prefix="dokiai_integrity_") as directory:
            root = Path(directory)
            output = root / "output.docx"
            before = root / "before.docx"
            document = Document()
            document.add_paragraph("封面")
            document.add_paragraph("摘要")
            document.add_paragraph("保留的论文正文")
            document.save(output)
            shutil.copyfile(output, before)
            with zipfile.ZipFile(before, "a") as package:
                package.writestr("word/media/removed-cover-image.bin", b"cover")

            result = validate_preservation(
                before,
                output,
                source_body_start=2,
                allow_front_matter=True,
            )

            self.assertTrue(result.passed, result.differences)


class FinalDeliveryPolicyTests(unittest.TestCase):
    def test_comments_are_removed_and_red_source_text_becomes_black(self) -> None:
        with tempfile.TemporaryDirectory(prefix="dokiai_finalizer_") as directory:
            path = Path(directory) / "reviewed.docx"
            document = Document()
            paragraph = document.add_paragraph()
            red_run = paragraph.add_run("必须保留的原稿文字")
            red_run.font.color.rgb = RGBColor(255, 0, 0)
            document.add_comment(red_run, text="仅用于核对格式", author="Reviewer")
            document.save(path)

            stats = finalize_docx(path)

            with zipfile.ZipFile(path) as package:
                names = package.namelist()
                self.assertFalse(any("comment" in name.lower() for name in names))
                document_xml = package.read("word/document.xml").decode("utf-8")
                self.assertNotIn("commentRange", document_xml)
                self.assertNotIn("commentReference", document_xml)
                self.assertIn('w:val="000000"', document_xml)
            self.assertGreater(stats["comment_parts_removed"], 0)
            self.assertGreater(stats["red_fonts_blackened"], 0)

    def test_red_template_instructions_are_not_copied_and_front_pages_become_sections(self) -> None:
        document = Document()
        document.add_paragraph("封面")
        second_page = document.add_paragraph("诚信声明")
        second_page.paragraph_format.page_break_before = True
        review = document.add_paragraph()
        review_run = review.add_run("此处红字只用于说明格式")
        review_run.font.color.rgb = RGBColor(255, 0, 0)
        body = document.add_heading("第一章 正文", level=1)
        body.paragraph_format.page_break_before = True

        removed = DocumentProcessor._remove_template_review_artifacts(document, 4)
        added = DocumentProcessor._isolate_front_matter_pages(document, 4)

        self.assertGreater(removed, 0)
        self.assertNotIn("此处红字只用于说明格式", "\n".join(p.text for p in document.paragraphs))
        self.assertGreaterEqual(added, 1)
        self.assertGreaterEqual(len(document.sections), 2)

    def test_each_numbered_chapter_starts_on_a_new_page(self) -> None:
        document = Document()
        first = document.add_heading("第一章 绪论", level=1)
        document.add_paragraph("正文")
        second = document.add_heading("第二章 方法", level=1)
        result = ProcessResult(Path("source.docx"), Path("output.docx"))

        DocumentProcessor._start_chapters_on_new_pages(document, 1, result)

        self.assertTrue(first.paragraph_format.page_break_before)
        self.assertTrue(second.paragraph_format.page_break_before)


if __name__ == "__main__":
    unittest.main()
